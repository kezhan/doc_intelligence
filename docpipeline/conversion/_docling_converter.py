"""
DoclingConverter — Conversion PDF vers DOCX via IBM Docling (ML open-source).

Porté depuis https://github.com/CHRISTMardochee/pdf2word et personnalisé.

C'est l'**alternative gratuite la plus proche d'Adobe** pour la conversion
PDF → DOCX éditable + visuellement fidèle :

- Modèles ML pré-entraînés par IBM Research (compréhension structurelle)
- Détection sémantique : titres, paragraphes, listes, tableaux, légendes
- Reading-order intelligent (multi-colonnes, encadrés)
- Reconnaissance de structure de tableaux (TableFormer)
- 100% offline après le 1er téléchargement des modèles (~500 Mo)

Pipeline :
  1. Docling ML → structure sémantique du document
  2. Construction du DOCX à partir du Markdown structuré généré
  3. Post-traitement PyMuPDF : application des polices, couleurs,
     tailles, fonds colorés extraits du PDF source

Prérequis :
  pip install docpipeline[docling]
  ou : pip install docling
"""

from __future__ import annotations

import logging
import os
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

logger = logging.getLogger(__name__)


class DoclingConverter:
    """
    Convertisseur PDF → DOCX via IBM Docling (ML).

    Aucun LLM (ce sont des modèles ML spécialisés document, pas génératifs).
    Aucune connexion réseau requise après le 1er run.
    """

    def __init__(self) -> None:
        try:
            from docling.document_converter import DocumentConverter  # type: ignore
            self._converter_class = DocumentConverter
        except ImportError as exc:
            raise ImportError(
                "DoclingConverter nécessite IBM Docling. Installer :\n"
                "    pip install docling\n"
                "ou : pip install docpipeline[docling]"
            ) from exc

    def convert(
        self,
        input_pdf:    str | Path,
        output_docx:  str | Path,
        pages:        list[int] | None = None,
    ) -> str:
        """
        Convertir via Docling ML + post-traitement visuel PyMuPDF.

        Input  : chemin PDF + chemin DOCX
        Output : chemin DOCX généré
        """
        import fitz

        input_pdf  = os.path.abspath(str(input_pdf))
        output_pdf = os.path.abspath(str(output_docx))

        if not os.path.isfile(input_pdf):
            raise FileNotFoundError(f"PDF introuvable : {input_pdf}")

        logger.info("DoclingConverter : %s -> %s", input_pdf, output_pdf)

        actual_pdf = input_pdf
        tmp_pdf    = None

        if pages is not None:
            tmp_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
            tmp_pdf.close()
            src = fitz.open(input_pdf)
            dst = fitz.open()
            try:
                for p in pages:
                    if p < len(src):
                        dst.insert_pdf(src, from_page=p, to_page=p)
                dst.save(tmp_pdf.name)
            finally:
                dst.close()
                src.close()
            actual_pdf = tmp_pdf.name

        try:
            # 1. Analyse ML par Docling
            logger.info("Analyse Docling ML (téléchargement modèles au 1er run)...")
            converter = self._converter_class()
            result    = converter.convert(actual_pdf)
            docling_doc = result.document

            # 2. Construction du DOCX depuis la sortie structurée Markdown
            doc = self._build_docx(docling_doc)

            # 3. Application des styles visuels du PDF source
            logger.info("Application des styles visuels (PyMuPDF)...")
            self._apply_pdf_visual_styles(doc, actual_pdf)

            for section in doc.sections:
                section.left_margin   = Cm(2.0)
                section.right_margin  = Cm(2.0)
                section.top_margin    = Cm(1.5)
                section.bottom_margin = Cm(1.5)

            doc.save(output_pdf)
            logger.info("DoclingConverter terminé : %s", output_pdf)
            return output_pdf

        finally:
            if tmp_pdf is not None:
                try:
                    os.unlink(tmp_pdf.name)
                except OSError:
                    pass

    # ── Construction DOCX depuis le Markdown structuré Docling ──────────────

    def _build_docx(self, docling_doc: Any) -> Document:
        doc                = Document()
        style              = doc.styles["Normal"]
        style.font.name    = "Calibri"
        style.font.size    = Pt(10)

        md_text = docling_doc.export_to_markdown()
        lines   = md_text.split("\n")

        i = 0
        while i < len(lines):
            line     = lines[i]
            stripped = line.strip()

            if not stripped or stripped == "<!-- image -->":
                i += 1
                continue

            # Titres H1/H2/H3
            if stripped.startswith("# ") and not stripped.startswith("## "):
                doc.add_heading(stripped[2:].strip(), level=1)
                i += 1; continue
            if stripped.startswith("## "):
                doc.add_heading(stripped[3:].strip(), level=2)
                i += 1; continue
            if stripped.startswith("### "):
                doc.add_heading(stripped[4:].strip(), level=3)
                i += 1; continue

            # Tableaux (format Markdown)
            if stripped.startswith("|"):
                table_lines = []
                while i < len(lines) and lines[i].strip().startswith("|"):
                    table_lines.append(lines[i].strip())
                    i += 1
                self._add_markdown_table(doc, table_lines)
                continue

            # Liste à puces
            if stripped.startswith("- ") or stripped.startswith("* "):
                para = doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after  = Pt(1)
                i += 1; continue

            # Liste numérotée
            if (len(stripped) > 2 and stripped[0].isdigit()
                    and stripped[1] in ".):"):
                para = doc.add_paragraph(stripped[2:].strip(), style="List Number")
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after  = Pt(1)
                i += 1; continue

            # Paragraphe normal
            para = doc.add_paragraph()
            run  = para.add_run(stripped)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            para.paragraph_format.space_before = Pt(2)
            para.paragraph_format.space_after  = Pt(2)
            i += 1

        return doc

    def _add_markdown_table(self, doc: Document, table_lines: list[str]) -> None:
        data_lines: list[list[str]] = []
        for line in table_lines:
            cells = [c.strip() for c in line.strip("|").split("|")]
            # Ignorer les lignes séparatrices |---|---|
            if all(c.replace("-", "").replace(":", "").strip() == "" for c in cells):
                continue
            if any(c for c in cells):
                data_lines.append(cells)

        if not data_lines:
            return

        max_cols = max(len(row) for row in data_lines)
        rows     = len(data_lines)
        if max_cols == 0 or rows == 0:
            return

        table = doc.add_table(rows=rows, cols=max_cols)
        table.autofit = True
        self._set_table_borders(table)

        for r, row_data in enumerate(data_lines):
            for c in range(min(len(row_data), max_cols)):
                cell = table.cell(r, c)
                txt  = row_data[c].strip()
                para = cell.paragraphs[0]
                if txt:
                    run = para.add_run(txt)
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    if r == 0:
                        run.font.bold       = True
                        run.font.color.rgb  = RGBColor(255, 255, 255)
                    if txt in ("●", "○", "◦", "•"):
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            if r == 0:
                for c in range(min(len(row_data), max_cols)):
                    cell = table.cell(0, c)
                    shd  = OxmlElement("w:shd")
                    shd.set(qn("w:fill"), "2E75B6")  # bleu corporate neutre
                    shd.set(qn("w:val"),  "clear")
                    cell._tc.get_or_add_tcPr().append(shd)

    # ── Post-traitement visuel via PyMuPDF ───────────────────────────────────

    def _apply_pdf_visual_styles(self, doc: Document, pdf_path: str) -> None:
        import fitz
        pdf_doc = fitz.open(pdf_path)
        try:
            visual = self._extract_visual_data(pdf_doc)
        finally:
            pdf_doc.close()

        self._apply_fonts_to_paragraphs(doc, visual)
        self._apply_heading_backgrounds(doc, visual)

    def _extract_visual_data(self, pdf_doc: Any) -> dict:
        import fitz
        data: dict = {"text_blocks": [], "backgrounds": [], "font_stats": {}}
        size_counts: dict[float, int] = {}

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]

            for block in page.get_text("dict").get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue
                        size      = round(span.get("size", 10), 1)
                        flags     = span.get("flags", 0)
                        color_int = span.get("color", 0)
                        is_bold   = bool(flags & (1 << 4))
                        is_italic = bool(flags & (1 << 1))
                        r = (color_int >> 16) & 0xFF
                        g = (color_int >>  8) & 0xFF
                        b = color_int & 0xFF
                        data["text_blocks"].append({
                            "page":   page_num,
                            "text":   text[:80],
                            "size":   size,
                            "bold":   is_bold,
                            "italic": is_italic,
                            "color":  f"{r:02X}{g:02X}{b:02X}",
                            "font":   span.get("font", ""),
                            "bbox":   span.get("bbox", []),
                        })
                        size_counts[size] = size_counts.get(size, 0) + len(text)

            for d in page.get_drawings():
                fill = d.get("fill")
                rect = d.get("rect")
                if fill is None or rect is None:
                    continue
                dr = fitz.Rect(rect)
                if dr.width < 50 or dr.height < 10:
                    continue
                hex_c = self._rgb_to_hex(fill)
                if hex_c == "FFFFFF":
                    continue
                data["backgrounds"].append({
                    "page":   page_num,
                    "rect":   rect,
                    "width":  dr.width,
                    "height": dr.height,
                    "color":  hex_c,
                    "y":      dr.y0,
                })

        data["font_stats"]["body_size"] = (
            max(size_counts, key=size_counts.get) if size_counts else 10.0
        )
        return data

    def _apply_fonts_to_paragraphs(self, doc: Document, visual: dict) -> None:
        body_size  = visual["font_stats"].get("body_size", 10.0)
        text_fonts = {}
        for block in visual["text_blocks"]:
            key = block["text"][:40].strip().lower()
            if key and key not in text_fonts:
                text_fonts[key] = block

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            match = text_fonts.get(text[:40].strip().lower())
            if not match:
                continue

            for run in para.runs:
                pdf_size = match["size"]
                run.font.size = Pt(pdf_size if pdf_size > body_size * 1.3 else body_size)
                if match["bold"]:
                    run.font.bold = True
                if match["italic"]:
                    run.font.italic = True
                color = match["color"]
                if color != "000000":
                    run.font.color.rgb = RGBColor(
                        int(color[0:2], 16),
                        int(color[2:4], 16),
                        int(color[4:6], 16),
                    )
                run.font.name = "Calibri"

    def _apply_heading_backgrounds(self, doc: Document, visual: dict) -> None:
        backgrounds = visual["backgrounds"]
        if not backgrounds:
            return

        header_bands = [bg for bg in backgrounds
                        if bg["width"] > 200 and 15 < bg["height"] < 50]
        info_boxes   = [bg for bg in backgrounds
                        if bg["width"] > 200 and bg["height"] > 50]

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            for block in visual["text_blocks"]:
                btext = block["text"].strip()
                if not btext or text[:30].lower() != btext[:30].lower():
                    continue
                bbox = block.get("bbox", [])
                if len(bbox) < 4:
                    continue
                ty, tp = bbox[1], block["page"]

                for band in header_bands:
                    if band["page"] != tp:
                        continue
                    br = band["rect"]
                    if (br.y0 - 5) <= ty <= (br.y1 + 5):
                        self._set_paragraph_shading(para, band["color"])
                        if self._is_dark_color(band["color"]):
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.bold      = True
                        break
                else:
                    for box in info_boxes:
                        if box["page"] != tp:
                            continue
                        br = box["rect"]
                        if (br.y0 - 5) <= ty <= (br.y1 + 5):
                            self._set_paragraph_shading(para, box["color"])
                            break
                break

    # ── Helpers ──────────────────────────────────────────────────────────────

    @staticmethod
    def _set_paragraph_shading(para: Any, hex_color: str) -> None:
        pPr = para._element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"),  "clear")
        pPr.append(shd)

    @staticmethod
    def _is_dark_color(hex_color: str) -> bool:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255
        return luminance < 0.5

    @staticmethod
    def _rgb_to_hex(rgb_tuple: tuple) -> str:
        r = int(min(max(rgb_tuple[0], 0), 1) * 255)
        g = int(min(max(rgb_tuple[1], 0), 1) * 255)
        b = int(min(max(rgb_tuple[2], 0), 1) * 255)
        return f"{r:02X}{g:02X}{b:02X}"

    def _set_table_borders(self, table: Any) -> None:
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        tbl  = table._tbl
        tbl_pr = tbl.find(f"{{{W_NS}}}tblPr")
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        existing = tbl_pr.find(f"{{{W_NS}}}tblBorders")
        if existing is not None:
            tbl_pr.remove(existing)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"),   "single")
            border.set(qn("w:sz"),    "4")
            border.set(qn("w:color"), "D0D0D0")
            border.set(qn("w:space"), "0")
            borders.append(border)
        tbl_pr.append(borders)
