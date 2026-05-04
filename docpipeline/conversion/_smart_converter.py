"""
SmartConverter — Conversion PDF vers DOCX via PyMuPDF.

Porté depuis https://github.com/CHRISTMardochee/pdf2word
et personnalisé pour le pipeline docpipeline.

Fonctionnalités :
- Préserve polices, tailles, couleurs, gras, italique
- Détecte les titres par taille de police (dynamique)
- Extrait et positionne les images
- Gère les mises en page multi-colonnes via tableaux Word
- Reconstruit les listes à puces/numérotées
- Sépare en-tête/pied de page du corps
- Détecte et restitue les tableaux avec couleurs de cellules
"""

from __future__ import annotations

import logging
import os
import re
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

_CONTROL_CHARS   = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]')
_BULLET_PATTERN  = re.compile(r'^[\s]*[•●○◦▪▸►–—\-\*]\s+')
_NUMBER_PATTERN  = re.compile(r'^[\s]*\d{1,3}[\.\)]\s+')
_W_NS            = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"

# Correspondance des noms de polices PDF → polices système
_FONT_MAP = {
    "TimesNewRoman":  "Times New Roman",
    "ArialMT":        "Arial",
    "CourierNew":     "Courier New",
    "Helvetica":      "Arial",
    "HelveticaNeue":  "Arial",
    "Times":          "Times New Roman",
    "Courier":        "Courier New",
}


class SmartConverter:
    """
    Convertisseur universel PDF → DOCX sans LLM.

    Utilise PyMuPDF pour une extraction précise du layout,
    des tableaux, des images et des styles de texte.
    """

    HEADER_Y_THRESHOLD = 0.06   # Zone en-tête : 6% supérieurs de la page
    FOOTER_Y_THRESHOLD = 0.92   # Zone pied de page : 8% inférieurs
    COLUMN_GAP_MIN     = 30.0   # Écart minimum entre colonnes (pts)
    MIN_IMAGE_SIZE     = 15     # Dimension minimale image à extraire (pts)

    def convert(
        self,
        pdf_path: str | Path,
        docx_path: str | Path,
        pages: list[int] | None = None,
    ) -> str:
        """
        Convertir un PDF en DOCX en préservant le layout complet.

        Input  : chemin PDF + chemin DOCX de sortie
        Output : chemin DOCX généré
        """
        pdf_path  = str(pdf_path)
        docx_path = str(docx_path)
        logger.info("SmartConverter: %s -> %s", pdf_path, docx_path)

        pdf_doc    = fitz.open(pdf_path)
        doc        = Document()
        style      = doc.styles["Normal"]
        style.font.name = "Calibri"
        style.font.size = Pt(10)

        font_stats   = self._analyze_fonts(pdf_doc, pages)
        page_numbers = pages if pages else list(range(len(pdf_doc)))

        with tempfile.TemporaryDirectory() as tmp_dir:
            for idx, page_num in enumerate(page_numbers):
                if page_num >= len(pdf_doc):
                    continue
                page = pdf_doc[page_num]
                logger.info("Page %d/%d", idx + 1, len(page_numbers))
                if idx > 0:
                    doc.add_page_break()
                self._process_page(doc, page, pdf_doc, tmp_dir, page_num, font_stats)

        pdf_doc.close()

        for section in doc.sections:
            section.left_margin   = Cm(2.0)
            section.right_margin  = Cm(2.0)
            section.top_margin    = Cm(1.5)
            section.bottom_margin = Cm(1.5)

        doc.save(docx_path)
        logger.info("SmartConverter terminé : %s", docx_path)
        return docx_path

    # ── Analyse des polices ───────────────────────────────────────────────────

    def _analyze_fonts(self, pdf_doc: fitz.Document, pages: list[int] | None) -> dict:
        """
        Détecter dynamiquement la taille corps + seuils titres H1/H2.
        Échantillonne jusqu'à 20 pages pour fiabilité.
        """
        size_counts: dict[float, int] = {}
        sample     = pages if pages else list(range(min(len(pdf_doc), 20)))

        for page_num in sample:
            if page_num >= len(pdf_doc):
                continue
            for block in pdf_doc[page_num].get_text("dict").get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue
                        size = round(span.get("size", 10), 1)
                        size_counts[size] = size_counts.get(size, 0) + len(text)

        if not size_counts:
            return {"body_size": 10.0, "heading1_min": 16.0, "heading2_min": 12.0}

        body_size    = max(size_counts, key=size_counts.get)
        sizes        = sorted(size_counts.keys())
        larger       = [s for s in sizes if s > body_size + 1.0]

        heading1_min = (larger[-1] if larger else body_size + 6.0)
        heading2_min = (larger[0]  if larger else body_size + 2.0)
        if len(larger) >= 2:
            heading1_min = max(heading1_min, heading2_min + 2)

        logger.info("Polices — corps=%.1f  H2≥%.1f  H1≥%.1f",
                    body_size, heading2_min, heading1_min)
        return {"body_size": body_size, "heading1_min": heading1_min,
                "heading2_min": heading2_min}

    # ── Traitement d'une page ─────────────────────────────────────────────────

    def _process_page(
        self,
        doc: Document,
        page: fitz.Page,
        pdf_doc: fitz.Document,
        tmp_dir: str,
        page_num: int,
        font_stats: dict,
    ) -> None:
        page_h   = page.rect.height
        page_w   = page.rect.width
        header_y = page_h * self.HEADER_Y_THRESHOLD
        footer_y = page_h * self.FOOTER_Y_THRESHOLD

        # Blocs texte dans le corps (hors en-tête/pied)
        text_dict   = page.get_text("dict")
        raw_blocks  = [b for b in text_dict.get("blocks", [])
                       if b["type"] == 0 and self._block_has_text(b)]
        body_blocks = [b for b in raw_blocks
                       if header_y <= b["bbox"][1] < footer_y]

        # Tableaux détectés
        tables_info = self._detect_tables(page)

        # Blocs hors tableaux
        non_table = [b for b in body_blocks
                     if not any(self._block_overlaps_table(b, t["bbox"]) for t in tables_info)]

        drawings = page.get_drawings()

        # Construction de la liste d'éléments triés par Y
        elements: list[tuple] = []

        for elem in self._collect_image_elements(page, pdf_doc, tmp_dir, page_num, header_y, footer_y):
            elements.append(elem)

        for t in tables_info:
            elements.append((t["bbox"][1], "table", t, drawings))

        columns = self._detect_columns(non_table, page_w)
        if columns == 2:
            mid        = page_w / 2.0
            full_width = [b for b in non_table if (b["bbox"][2] - b["bbox"][0]) > page_w * 0.55]
            left_col   = [b for b in non_table if b["bbox"][2] < mid + 20
                          and (b["bbox"][2] - b["bbox"][0]) <= page_w * 0.55]
            right_col  = [b for b in non_table if b["bbox"][0] > mid - 20
                          and (b["bbox"][2] - b["bbox"][0]) <= page_w * 0.55]
            for b in full_width:
                elements.append((b["bbox"][1], "text_block", b))
            if left_col or right_col:
                min_y = min(
                    left_col[0]["bbox"][1]  if left_col  else 9999,
                    right_col[0]["bbox"][1] if right_col else 9999,
                )
                elements.append((min_y, "columns", left_col, right_col))
        else:
            for b in non_table:
                elements.append((b["bbox"][1], "text_block", b))

        elements.sort(key=lambda x: x[0])

        for elem in elements:
            kind = elem[1]
            if kind == "table":
                self._render_table(doc, elem[2], elem[3], font_stats)
            elif kind == "text_block":
                self._add_block(doc, elem[2], font_stats)
            elif kind == "columns":
                left  = sorted(elem[2], key=lambda b: b["bbox"][1])
                right = sorted(elem[3], key=lambda b: b["bbox"][1])
                self._add_column_table(doc, left, right, font_stats)
            elif kind == "image":
                self._add_image_element(doc, elem[2])

    # ── Détection des tableaux ────────────────────────────────────────────────

    def _detect_tables(self, page: fitz.Page) -> list[dict]:
        try:
            result = page.find_tables()
            tables = []
            for t in result.tables:
                data = t.extract()
                if not data or len(data) < 2:
                    continue
                tables.append({
                    "bbox":      t.bbox,
                    "rows":      data,
                    "row_count": len(data),
                    "col_count": t.col_count,
                })
            return tables
        except Exception as exc:
            logger.warning("Détection tableaux échouée : %s", str(exc)[:80])
            return []

    def _block_overlaps_table(self, block: dict, table_bbox: tuple) -> bool:
        bx0, by0, bx1, by1 = block["bbox"]
        tx0, ty0, tx1, ty1 = table_bbox
        ox0, oy0 = max(bx0, tx0), max(by0, ty0)
        ox1, oy1 = min(bx1, tx1), min(by1, ty1)
        if ox0 >= ox1 or oy0 >= oy1:
            return False
        block_area   = max((bx1 - bx0) * (by1 - by0), 1)
        overlap_area = (ox1 - ox0) * (oy1 - oy0)
        return (overlap_area / block_area) > 0.5

    # ── Rendu des tableaux ────────────────────────────────────────────────────

    def _render_table(self, doc: Document, info: dict, drawings: list, font_stats: dict) -> None:
        rows_data  = info["rows"]
        row_count  = info["row_count"]
        col_count  = info["col_count"]
        table_bbox = info["bbox"]

        table = doc.add_table(rows=row_count, cols=col_count)
        table.autofit = True
        self._set_table_borders(table, show=True)

        cell_colors = self._get_cell_colors(table_bbox, row_count, col_count, rows_data, drawings)

        for r, row_data in enumerate(rows_data):
            for c in range(min(len(row_data), col_count)):
                cell      = table.cell(r, c)
                cell_text = self._sanitize(row_data[c] or "").strip()
                para      = cell.paragraphs[0]
                if cell_text:
                    run = para.add_run(cell_text)
                    run.font.size = Pt(font_stats["body_size"])
                    run.font.name = "Calibri"
                    if r == 0:
                        run.font.bold         = True
                        run.font.color.rgb    = RGBColor(255, 255, 255)
                    if cell_text.strip() in ("●", "○", "◦", "•"):
                        para.alignment = WD_ALIGN_PARAGRAPH.CENTER

                color_key = (r, c)
                if color_key in cell_colors:
                    self._set_cell_shading(cell, cell_colors[color_key])
                elif r == 0:
                    self._set_cell_shading(cell, "2E75B6")  # Bleu corporate neutre

    def _get_cell_colors(
        self,
        table_bbox: tuple,
        row_count: int,
        col_count: int,
        rows_data: list,
        drawings: list,
    ) -> dict:
        colors: dict[tuple, str] = {}
        tx0, ty0, tx1, ty1 = table_bbox
        col_w  = (tx1 - tx0) / max(col_count, 1)
        row_h  = (ty1 - ty0) / max(row_count, 1)

        filled = [
            (fitz.Rect(d["rect"]), d["fill"])
            for d in drawings
            if d.get("fill") and d.get("rect")
            and fitz.Rect(d["rect"]).intersects(fitz.Rect(tx0, ty0, tx1, ty1))
            and fitz.Rect(d["rect"]).width > 10
            and fitz.Rect(d["rect"]).height > 10
        ]

        for r in range(row_count):
            for c in range(col_count):
                cx = tx0 + c * col_w + col_w / 2
                cy = ty0 + r * row_h + row_h / 2
                for dr, fill in filled:
                    if dr.contains(fitz.Point(cx, cy)):
                        hex_c = self._rgb_to_hex(fill)
                        if hex_c != "FFFFFF":
                            colors[(r, c)] = hex_c
                        break
        return colors

    @staticmethod
    def _rgb_to_hex(rgb: tuple) -> str:
        r = int(min(max(rgb[0], 0), 1) * 255)
        g = int(min(max(rgb[1], 0), 1) * 255)
        b = int(min(max(rgb[2], 0), 1) * 255)
        return f"{r:02X}{g:02X}{b:02X}"

    @staticmethod
    def _set_cell_shading(cell, hex_color: str) -> None:
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"),  "clear")
        cell._tc.get_or_add_tcPr().append(shd)

    # ── Images ────────────────────────────────────────────────────────────────

    def _collect_image_elements(
        self,
        page: fitz.Page,
        pdf_doc: fitz.Document,
        tmp_dir: str,
        page_num: int,
        header_y: float,
        footer_y: float,
    ) -> list[tuple]:
        elements: list[tuple] = []
        seen: set[int] = set()

        for info in page.get_image_info(xrefs=True):
            xref = info.get("xref", 0)
            if xref in seen or xref == 0:
                continue
            seen.add(xref)

            bbox = info.get("bbox", [])
            if not bbox or len(bbox) < 4:
                continue

            w_pt, h_pt = abs(bbox[2] - bbox[0]), abs(bbox[3] - bbox[1])
            if w_pt < self.MIN_IMAGE_SIZE or h_pt < self.MIN_IMAGE_SIZE:
                continue
            if not (header_y <= bbox[1] < footer_y):
                continue

            img_path = os.path.join(tmp_dir, f"img_p{page_num}_x{xref}.png")
            try:
                clip  = fitz.Rect(bbox) & page.rect
                if clip.is_empty or clip.width < 5 or clip.height < 5:
                    continue
                scale = min(3.0, max(2.0, 600.0 / max(clip.width, clip.height)))
                page.get_pixmap(matrix=fitz.Matrix(scale, scale), clip=clip, alpha=False).save(img_path)
            except Exception as exc:
                logger.warning("Image xref=%d : %s", xref, str(exc)[:80])
                continue

            if not os.path.isfile(img_path):
                continue

            max_w_in = (page.rect.width - 80) / 72.0
            width_in = min(w_pt / 72.0, max_w_in)
            if width_in < 0.3:
                width_in = min(2.0, max_w_in)

            elements.append((bbox[1], "image", {"path": img_path, "width_in": width_in, "xref": xref}))
        return elements

    def _add_image_element(self, doc: Document, img_data: dict) -> None:
        try:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.add_run().add_picture(img_data["path"], width=Inches(img_data["width_in"]))
            para.paragraph_format.space_before = Pt(4)
            para.paragraph_format.space_after  = Pt(4)
        except Exception as exc:
            logger.warning("Image insert xref=%d : %s", img_data["xref"], str(exc)[:80])

    # ── Colonnes ──────────────────────────────────────────────────────────────

    def _detect_columns(self, blocks: list, page_width: float) -> int:
        if not blocks:
            return 1
        mid        = page_width / 2.0
        left_only  = sum(1 for b in blocks if b["bbox"][2] < mid + 20
                         and (b["bbox"][2] - b["bbox"][0]) <= page_width * 0.55)
        right_only = sum(1 for b in blocks if b["bbox"][0] > mid - 20
                         and (b["bbox"][2] - b["bbox"][0]) <= page_width * 0.55)
        return 2 if left_only >= 3 and right_only >= 3 else 1

    def _add_column_table(
        self,
        doc: Document,
        left_blocks: list,
        right_blocks: list,
        font_stats: dict,
    ) -> None:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        self._set_table_borders(table, show=False)

        for i, block in enumerate(left_blocks):
            para = table.cell(0, 0).paragraphs[0] if i == 0 else table.cell(0, 0).add_paragraph()
            self._fill_paragraph(para, block, font_stats)

        for i, block in enumerate(right_blocks):
            para = table.cell(0, 1).paragraphs[0] if i == 0 else table.cell(0, 1).add_paragraph()
            self._fill_paragraph(para, block, font_stats)

    # ── Blocs texte ───────────────────────────────────────────────────────────

    def _add_block(self, doc: Document, block: dict, font_stats: dict) -> None:
        text       = self._get_block_text(block).strip()
        style_type = self._classify_block(block, font_stats)
        is_bullet  = bool(_BULLET_PATTERN.match(text))
        is_numbered = bool(_NUMBER_PATTERN.match(text))

        if style_type == "heading1":
            para = doc.add_heading(level=1)
        elif style_type == "heading2":
            para = doc.add_heading(level=2)
        elif is_bullet:
            para = doc.add_paragraph(style="List Bullet")
        elif is_numbered:
            para = doc.add_paragraph(style="List Number")
        else:
            para = doc.add_paragraph()

        self._fill_paragraph(para, block, font_stats)

    def _classify_block(self, block: dict, font_stats: dict) -> str:
        max_size = 0
        is_bold  = False
        for line in block.get("lines", []):
            for span in line.get("spans", []):
                max_size = max(max_size, span.get("size", 0))
                if span.get("flags", 0) & 16:
                    is_bold = True

        if max_size >= font_stats["heading1_min"]:
            return "heading1"
        if max_size >= font_stats["heading2_min"] and is_bold:
            return "heading2"
        if max_size > font_stats["body_size"] + 1.5 and is_bold:
            return "heading2"
        return "body"

    def _fill_paragraph(self, para, block: dict, font_stats: dict) -> None:
        for run in para.runs:
            run.text = ""

        first = True
        for line_idx, line in enumerate(block.get("lines", [])):
            if line_idx > 0 and not first:
                prev = para.runs[-1].text if para.runs else ""
                if prev.endswith("-"):
                    para.runs[-1].text = prev[:-1]
                else:
                    para.add_run(" ")

            for span in line.get("spans", []):
                text = self._sanitize(span.get("text", ""))
                if not text:
                    continue
                first = False
                run   = para.add_run(text)

                size  = span.get("size", font_stats["body_size"])
                run.font.size = Pt(size)

                ci = span.get("color", 0)
                run.font.color.rgb = RGBColor((ci >> 16) & 0xFF, (ci >> 8) & 0xFF, ci & 0xFF)

                flags = span.get("flags", 0)
                if flags & 16:
                    run.font.bold   = True
                if flags & 2:
                    run.font.italic = True

                clean = self._clean_font_name(span.get("font", ""))
                if clean:
                    run.font.name = clean

        pf         = para.paragraph_format
        style_type = self._classify_block(block, font_stats)
        if style_type == "heading1":
            pf.space_before = Pt(16); pf.space_after = Pt(6)
        elif style_type == "heading2":
            pf.space_before = Pt(10); pf.space_after = Pt(4)
        else:
            pf.space_before = Pt(2);  pf.space_after = Pt(2)
            pf.line_spacing = Pt(14)

    # ── Bordures tableaux ─────────────────────────────────────────────────────

    def _set_table_borders(self, table, *, show: bool) -> None:
        tbl    = table._tbl
        tbl_pr = tbl.find(f"{{{_W_NS}}}tblPr")
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)

        existing = tbl_pr.find(f"{{{_W_NS}}}tblBorders")
        if existing is not None:
            tbl_pr.remove(existing)

        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{edge}")
            if show:
                border.set(qn("w:val"),   "single")
                border.set(qn("w:sz"),    "4")
                border.set(qn("w:color"), "D0D0D0")
            else:
                border.set(qn("w:val"),   "none")
                border.set(qn("w:sz"),    "0")
                border.set(qn("w:color"), "FFFFFF")
            border.set(qn("w:space"), "0")
            borders.append(border)
        tbl_pr.append(borders)

    # ── Utilitaires ───────────────────────────────────────────────────────────

    @staticmethod
    def _sanitize(text: str) -> str:
        return _CONTROL_CHARS.sub("", text)

    @staticmethod
    def _clean_font_name(font_name: str) -> str:
        if not font_name:
            return ""
        if "+" in font_name:
            font_name = font_name.split("+", 1)[-1]
        base = font_name.split("-")[0].split(",")[0]
        return _FONT_MAP.get(base, base)

    def _block_has_text(self, block: dict) -> bool:
        return any(
            self._sanitize(span.get("text", "")).strip()
            for line in block.get("lines", [])
            for span in line.get("spans", [])
        )

    def _get_block_text(self, block: dict) -> str:
        return " ".join(
            self._sanitize(span.get("text", ""))
            for line in block.get("lines", [])
            for span in line.get("spans", [])
            if self._sanitize(span.get("text", ""))
        )
