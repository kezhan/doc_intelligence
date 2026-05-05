"""
OverlayConverter — Conversion PDF → DOCX avec fidélité visuelle ET édition libre.

Approche unique : combine le meilleur des deux mondes.

Principe :
  Pour chaque page du PDF :
    1. Rendu de la page entière en image haute résolution → arrière-plan
    2. Pour chaque span de texte du PDF :
       - Création d'une zone de texte Word (text box) flottante
       - Positionnée pixel-précis au-dessus de l'image
       - Avec la police, taille, couleur d'origine
       - Texte invisible (couleur = fond, ou via un overlay)
         OU texte visible si on veut masquer celui de l'image

Résultat :
  - L'utilisateur voit le PDF original (image en arrière-plan)
  - Cliquer sur n'importe quel texte → édition libre
  - Mise en page strictement préservée (positions absolues)

Aucun LLM, aucun cloud, aucun compte requis.
Cette approche est ce que fait Acrobat Pro DC en mode édition.
"""

from __future__ import annotations

import logging
import os
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.shape import WD_INLINE_SHAPE
from docx.oxml import OxmlElement
from docx.oxml.ns import nsmap, qn
from docx.shared import Cm, Emu, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

# Namespaces Office
_W_NS  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_WP_NS = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
_A_NS  = "http://schemas.openxmlformats.org/drawingml/2006/main"
_R_NS  = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
_PIC_NS = "http://schemas.openxmlformats.org/drawingml/2006/picture"
_MC_NS = "http://schemas.openxmlformats.org/markup-compatibility/2006"
_WPS_NS = "http://schemas.microsoft.com/office/word/2010/wordprocessingShape"
_W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"

EMU_PER_PT = 12700  # 1 pt = 12700 EMU
EMU_PER_IN = 914400


class OverlayConverter:
    """
    Convertisseur PDF → DOCX hybride avancé : image fidèle + texte éditable
    positionné précisément au-dessus.

    Aucun LLM. Aucun cloud. 100% offline.
    """

    def __init__(
        self,
        *,
        dpi: int = 200,
        text_visibility: str = "transparent",  # 'transparent' | 'visible'
    ) -> None:
        """
        Args:
            dpi              : résolution rendu pages (200 par défaut, 300 = print)
            text_visibility  : 'transparent' (texte sur l'image) ou 'visible'
                               (texte coloré, l'image est masquée par les boxes)
        """
        self.dpi = dpi
        self.text_visibility = text_visibility

    def convert(
        self,
        input_pdf:    str | Path,
        output_docx:  str | Path,
        pages:        list[int] | None = None,
    ) -> str:
        """
        Convertir un PDF en DOCX avec image de fond + zones texte éditables.

        Input  : chemin PDF + chemin DOCX
        Output : chemin DOCX généré
        """
        input_pdf  = str(input_pdf)
        output_pdf = str(output_docx)

        if not os.path.isfile(input_pdf):
            raise FileNotFoundError(f"PDF introuvable : {input_pdf}")

        logger.info("OverlayConverter (dpi=%d) : %s -> %s",
                    self.dpi, input_pdf, output_pdf)

        pdf_doc = fitz.open(input_pdf)
        doc     = Document()

        # Supprimer le paragraphe vide initial
        if doc.paragraphs:
            elem = doc.paragraphs[0]._element
            elem.getparent().remove(elem)

        page_numbers = pages if pages else list(range(len(pdf_doc)))

        with tempfile.TemporaryDirectory() as tmp_dir:
            for idx, page_num in enumerate(page_numbers):
                if page_num >= len(pdf_doc):
                    continue

                page = pdf_doc[page_num]
                logger.info("Page %d/%d", idx + 1, len(page_numbers))

                # 1. Section avec dimensions exactes du PDF
                section = doc.sections[0] if idx == 0 else doc.add_section()
                w_in = page.rect.width  / 72.0
                h_in = page.rect.height / 72.0
                section.page_width    = Inches(w_in)
                section.page_height   = Inches(h_in)
                section.left_margin   = Cm(0)
                section.right_margin  = Cm(0)
                section.top_margin    = Cm(0)
                section.bottom_margin = Cm(0)
                section.orientation   = (WD_ORIENT.LANDSCAPE if w_in > h_in
                                         else WD_ORIENT.PORTRAIT)

                # 2. Rendu page en image (arrière-plan) — float behind text
                img_path = os.path.join(tmp_dir, f"page_{page_num}.png")
                self._render_page(page, img_path)
                self._add_background_image(doc, img_path, page.rect, page_num)

                # 3. Extraction des spans de texte avec positions
                spans = self._extract_text_spans(page)
                logger.info("  → %d spans texte à overlayer", len(spans))

                # 4. Insertion des spans en zones de texte flottantes
                for span_idx, span in enumerate(spans):
                    self._add_text_overlay(doc, span, page_num, span_idx)

        pdf_doc.close()
        doc.save(output_pdf)
        logger.info("OverlayConverter terminé : %s", output_pdf)
        return output_pdf

    # ── Rendu page ────────────────────────────────────────────────────────────

    def _render_page(self, page: fitz.Page, output_path: str) -> None:
        zoom = self.dpi / 72.0
        pix  = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        pix.save(output_path)

    # ── Extraction spans de texte avec position ───────────────────────────────

    def _extract_text_spans(self, page: fitz.Page) -> list[dict]:
        spans: list[dict] = []
        text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                for span in line.get("spans", []):
                    text = self._sanitize(span.get("text", ""))
                    if not text.strip():
                        continue
                    bbox = span.get("bbox", [0, 0, 0, 0])
                    color_int = span.get("color", 0)
                    flags     = span.get("flags", 0)
                    spans.append({
                        "text":   text,
                        "x0":     bbox[0],
                        "y0":     bbox[1],
                        "x1":     bbox[2],
                        "y1":     bbox[3],
                        "size":   span.get("size", 10),
                        "font":   span.get("font", ""),
                        "bold":   bool(flags & (1 << 4)),
                        "italic": bool(flags & (1 << 1)),
                        "color":  color_int,
                    })
        return spans

    # ── Image de fond (anchored, behind text) ─────────────────────────────────

    def _add_background_image(
        self,
        doc: Document,
        img_path: str,
        page_rect: fitz.Rect,
        page_num: int,
    ) -> None:
        """Insère l'image de la page en arrière-plan ancré (behind text)."""
        # Paragraphe support pour l'ancrage
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        para.paragraph_format.line_spacing = 1.0
        run  = para.add_run()

        # Insertion image inline puis transformation en anchored behind
        inline_shape = run.add_picture(img_path,
                                        width=Inches(page_rect.width / 72.0))

        # Récupérer l'élément <w:drawing> et transformer inline → anchor
        drawing = run._element.findall(qn("w:drawing"))
        if not drawing:
            return
        drawing_el = drawing[0]
        inline_el  = drawing_el.find(qn("wp:inline"))
        if inline_el is None:
            # Compatibilité avec différents schémas
            for child in drawing_el:
                if child.tag.endswith("inline"):
                    inline_el = child
                    break
        if inline_el is None:
            return

        # Construire un <wp:anchor> "behind text"
        page_w_emu = int(page_rect.width  / 72.0 * EMU_PER_IN)
        page_h_emu = int(page_rect.height / 72.0 * EMU_PER_IN)

        anchor = OxmlElement("wp:anchor")
        anchor.set("distT", "0")
        anchor.set("distB", "0")
        anchor.set("distL", "0")
        anchor.set("distR", "0")
        anchor.set("simplePos", "0")
        anchor.set("relativeHeight", "1")
        anchor.set("behindDoc", "1")  # ⭐ derrière le texte
        anchor.set("locked", "0")
        anchor.set("layoutInCell", "1")
        anchor.set("allowOverlap", "1")

        # simplePos (requis même si simplePos=0)
        sp = OxmlElement("wp:simplePos")
        sp.set("x", "0"); sp.set("y", "0")
        anchor.append(sp)

        # Position H/V relatives à la page
        ph = OxmlElement("wp:positionH")
        ph.set("relativeFrom", "page")
        ph_off = OxmlElement("wp:posOffset")
        ph_off.text = "0"
        ph.append(ph_off)
        anchor.append(ph)

        pv = OxmlElement("wp:positionV")
        pv.set("relativeFrom", "page")
        pv_off = OxmlElement("wp:posOffset")
        pv_off.text = "0"
        pv.append(pv_off)
        anchor.append(pv)

        # Extent (taille image = taille page)
        extent = OxmlElement("wp:extent")
        extent.set("cx", str(page_w_emu))
        extent.set("cy", str(page_h_emu))
        anchor.append(extent)

        # effectExtent
        ee = OxmlElement("wp:effectExtent")
        ee.set("l", "0"); ee.set("t", "0"); ee.set("r", "0"); ee.set("b", "0")
        anchor.append(ee)

        # wrapNone (pas de wrap, l'image est libre)
        anchor.append(OxmlElement("wp:wrapNone"))

        # Copier docPr et graphic depuis l'inline
        for tag_name in ("wp:docPr", "wp:cNvGraphicFramePr", "a:graphic"):
            child = self._find_first(inline_el, tag_name)
            if child is not None:
                anchor.append(child)

        # Mettre à jour la taille du graphic xfrm si présent
        for ext in anchor.findall(f".//{{{_A_NS}}}ext"):
            ext.set("cx", str(page_w_emu))
            ext.set("cy", str(page_h_emu))

        # Remplacer inline par anchor
        drawing_el.remove(inline_el)
        drawing_el.append(anchor)

    # ── Zone de texte flottante par span ──────────────────────────────────────

    def _add_text_overlay(
        self,
        doc: Document,
        span: dict,
        page_num: int,
        span_idx: int,
    ) -> None:
        """Insère une zone de texte Word positionnée précisément au-dessus de l'image."""
        # Convertir les coordonnées PDF (en points) → EMU
        x_emu = int(span["x0"] * EMU_PER_PT)
        y_emu = int(span["y0"] * EMU_PER_PT)
        w_emu = max(int((span["x1"] - span["x0"]) * EMU_PER_PT), int(0.5 * EMU_PER_IN))
        h_emu = max(int((span["y1"] - span["y0"]) * EMU_PER_PT), int(span["size"] * EMU_PER_PT * 1.2))

        # Couleur du texte
        color_int = span["color"]
        r = (color_int >> 16) & 0xFF
        g = (color_int >>  8) & 0xFF
        b = color_int & 0xFF

        if self.text_visibility == "transparent":
            # Texte invisible (couleur très claire) — l'image fournit le visuel
            text_color_hex = "FEFEFE"
        else:
            text_color_hex = f"{r:02X}{g:02X}{b:02X}"

        # Construire le bloc XML w:p contenant w:r > mc:AlternateContent > wps:wsp
        # (text box flottant ancré sur la page)
        text_xml = self._build_text_box_xml(
            x_emu=x_emu, y_emu=y_emu, w_emu=w_emu, h_emu=h_emu,
            text=span["text"],
            font_size_pt=span["size"],
            bold=span["bold"],
            italic=span["italic"],
            color_hex=text_color_hex,
            font_name=self._clean_font(span["font"]),
            unique_id=page_num * 100000 + span_idx,
        )

        # Ajouter à la fin du body
        from lxml import etree
        body = doc.element.body
        new_p = etree.fromstring(text_xml)
        body.append(new_p)

    def _build_text_box_xml(
        self, *, x_emu: int, y_emu: int, w_emu: int, h_emu: int,
        text: str, font_size_pt: float, bold: bool, italic: bool,
        color_hex: str, font_name: str, unique_id: int,
    ) -> str:
        """Génère le XML OOXML pour une zone de texte flottante ancrée."""
        # Échapper le texte pour XML
        text_esc = (
            text.replace("&", "&amp;")
                .replace("<", "&lt;")
                .replace(">", "&gt;")
        )
        # Demi-pt pour w:sz
        sz_half_pt = int(font_size_pt * 2)

        bold_xml   = '<w:b/>'   if bold   else ''
        italic_xml = '<w:i/>'   if italic else ''

        return f'''<w:p xmlns:w="{_W_NS}" xmlns:wp="{_WP_NS}" xmlns:a="{_A_NS}" xmlns:r="{_R_NS}" xmlns:wps="{_WPS_NS}" xmlns:mc="{_MC_NS}" xmlns:w14="{_W14_NS}" xmlns:pic="{_PIC_NS}">
  <w:pPr><w:spacing w:before="0" w:after="0"/></w:pPr>
  <w:r>
    <w:rPr><w:noProof/></w:rPr>
    <mc:AlternateContent>
      <mc:Choice Requires="wps">
        <w:drawing>
          <wp:anchor distT="0" distB="0" distL="0" distR="0" simplePos="0" relativeHeight="{2 + unique_id}" behindDoc="0" locked="0" layoutInCell="1" allowOverlap="1">
            <wp:simplePos x="0" y="0"/>
            <wp:positionH relativeFrom="page"><wp:posOffset>{x_emu}</wp:posOffset></wp:positionH>
            <wp:positionV relativeFrom="page"><wp:posOffset>{y_emu}</wp:posOffset></wp:positionV>
            <wp:extent cx="{w_emu}" cy="{h_emu}"/>
            <wp:effectExtent l="0" t="0" r="0" b="0"/>
            <wp:wrapNone/>
            <wp:docPr id="{1000 + unique_id}" name="TextBox{unique_id}"/>
            <wp:cNvGraphicFramePr/>
            <a:graphic>
              <a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">
                <wps:wsp>
                  <wps:cNvSpPr txBox="1"/>
                  <wps:spPr>
                    <a:xfrm>
                      <a:off x="0" y="0"/>
                      <a:ext cx="{w_emu}" cy="{h_emu}"/>
                    </a:xfrm>
                    <a:prstGeom prst="rect"><a:avLst/></a:prstGeom>
                    <a:noFill/>
                    <a:ln><a:noFill/></a:ln>
                  </wps:spPr>
                  <wps:txbx>
                    <w:txbxContent>
                      <w:p>
                        <w:pPr>
                          <w:spacing w:before="0" w:after="0" w:line="240" w:lineRule="auto"/>
                          <w:ind w:left="0" w:right="0"/>
                        </w:pPr>
                        <w:r>
                          <w:rPr>
                            <w:rFonts w:ascii="{font_name}" w:hAnsi="{font_name}"/>
                            <w:color w:val="{color_hex}"/>
                            <w:sz w:val="{sz_half_pt}"/>
                            {bold_xml}{italic_xml}
                          </w:rPr>
                          <w:t xml:space="preserve">{text_esc}</w:t>
                        </w:r>
                      </w:p>
                    </w:txbxContent>
                  </wps:txbx>
                  <wps:bodyPr rot="0" spcFirstLastPara="0" vertOverflow="visible" horzOverflow="visible" wrap="none" lIns="0" tIns="0" rIns="0" bIns="0" anchor="t"/>
                </wps:wsp>
              </a:graphicData>
            </a:graphic>
          </wp:anchor>
        </w:drawing>
      </mc:Choice>
      <mc:Fallback>
        <w:pict>
          <w:r><w:t>{text_esc}</w:t></w:r>
        </w:pict>
      </mc:Fallback>
    </mc:AlternateContent>
  </w:r>
</w:p>'''

    # ── Helpers ──────────────────────────────────────────────────────────────

    @staticmethod
    def _sanitize(text: str) -> str:
        import re
        return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)

    @staticmethod
    def _clean_font(font_name: str) -> str:
        if not font_name:
            return "Arial"
        if "+" in font_name:
            font_name = font_name.split("+", 1)[-1]
        base = font_name.split("-")[0].split(",")[0]
        font_map = {
            "TimesNewRoman": "Times New Roman", "ArialMT": "Arial",
            "CourierNew": "Courier New", "Helvetica": "Arial",
            "HelveticaNeue": "Arial", "Times": "Times New Roman",
            "Courier": "Courier New",
        }
        return font_map.get(base, base or "Arial")

    @staticmethod
    def _find_first(parent, qname_with_prefix: str):
        """Trouve le 1er enfant matchant 'prefix:tag' (avec namespace)."""
        prefix, _, local = qname_with_prefix.partition(":")
        ns_map = {
            "wp": _WP_NS, "w": _W_NS, "a": _A_NS, "r": _R_NS, "pic": _PIC_NS,
        }
        ns = ns_map.get(prefix)
        if not ns:
            return None
        return parent.find(f"{{{ns}}}{local}")
