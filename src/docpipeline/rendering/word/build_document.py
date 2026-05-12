"""
build_document.py — Reconstruire un .docx a partir d'un runs_df modifie.

Symetrique de `parsing/word/parse_word.py`. Step 2 du build order Tome 2
translation (cf. CLAUDE_tome2_translation.md Sec.1.5) :

    extract  : parse_word(source)['span_df']  -> runs avec span_id stable
    modify   : on remplace seulement le 'text' de chaque run
    rebuild  : build_word_document(translated_runs_df, source, output)
               -> ouvre source comme template, walk paragraphs + tables,
                  remplace .text par span_id, save

Le `span_id` est la cle stable :
  - hors table : `w_<para>_<run>`
  - dans table : `w_t_<table>_<row>_<col>_<para>_<run>`

Aucun LLM (regle CLAUDE.md : LLM reserve a translation/summarization/
Excel SQL agent ; le rendering est pure manipulation python-docx).
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document


def build_word_document(
    translated_runs_df: pd.DataFrame,
    source_path,
    output_path,
) -> dict:
    """
    Reconstruit un .docx en remplacant le texte de certains runs, en
    preservant tous les styles d'origine (font, size, bold, italic, color,
    underline, highlight, etc.).

    Args:
        translated_runs_df : DataFrame avec colonnes 'span_id' et 'text'.
                              Les runs non listes gardent leur texte original.
                              span_id formats :
                                - `w_<para>_<run>` (hors table)
                                - `w_t_<table>_<row>_<col>_<para>_<run>` (cell)
        source_path        : chemin du .docx source (utilise comme template)
        output_path        : chemin du .docx de sortie

    Returns:
        dict {output_path, runs_replaced, runs_unchanged, runs_skipped, warnings}
    """
    source_path = Path(source_path)
    output_path = Path(output_path)

    if "span_id" not in translated_runs_df.columns or "text" not in translated_runs_df.columns:
        raise ValueError(
            "translated_runs_df doit contenir les colonnes 'span_id' et 'text'."
        )

    runs_by_span_id: dict[str, str] = {
        row["span_id"]: row["text"]
        for _, row in translated_runs_df.iterrows()
    }

    doc = Document(str(source_path))

    replaced = 0
    unchanged = 0
    skipped: list[str] = []
    warnings: list[str] = []

    # 1) Walk top-level paragraphs (hors tables)
    for para_idx, para in enumerate(doc.paragraphs):
        for run_idx, run in enumerate(para.runs):
            span_id = f"w_{para_idx}_{run_idx}"
            if span_id in runs_by_span_id:
                new_text = runs_by_span_id[span_id]
                if new_text != run.text:
                    run.text = new_text
                    replaced += 1
                else:
                    unchanged += 1
            else:
                skipped.append(span_id)

    # 2) Walk tables : doc.paragraphs n'inclut PAS les paragraphs des cells
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for cell_para_idx, cell_para in enumerate(cell.paragraphs):
                    for cell_run_idx, cell_run in enumerate(cell_para.runs):
                        span_id = (
                            f"w_t_{table_idx}_{row_idx}_{col_idx}"
                            f"_{cell_para_idx}_{cell_run_idx}"
                        )
                        if span_id in runs_by_span_id:
                            new_text = runs_by_span_id[span_id]
                            if new_text != cell_run.text:
                                cell_run.text = new_text
                                replaced += 1
                            else:
                                unchanged += 1
                        else:
                            skipped.append(span_id)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    return {
        "output_path":     output_path,
        "runs_replaced":   replaced,
        "runs_unchanged":  unchanged,
        "runs_skipped":    len(skipped),
        "warnings":        warnings,
    }


if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) != 3:
        print("Usage: python build_document.py <source.docx> <output.docx>", file=sys.stderr)
        print("       (round-trip identite : reconstruit source a output sans modif)", file=sys.stderr)
        sys.exit(1)

    src = Path(sys.argv[1])
    out = Path(sys.argv[2])

    from docpipeline.parsing.word.parse_word import parse_word
    runs_df = parse_word(src)["span_df"]
    result = build_word_document(runs_df, src, out)
    print(json.dumps({k: str(v) if isinstance(v, Path) else v
                      for k, v in result.items()}, indent=2, ensure_ascii=False))
