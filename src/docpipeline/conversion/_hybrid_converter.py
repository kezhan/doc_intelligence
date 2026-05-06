"""
HybridConverter — Conversion PDF vers DOCX en mode "image fidèle + texte invisible".

Porté depuis https://github.com/CHRISTMardochee/pdf2word et personnalisé.

Principe :
  - Chaque page du PDF est rendue en image haute résolution
  - L'image est insérée dans le DOCX dans une section aux dimensions exactes
  - Le texte natif est extrait et superposé en blanc invisible (1pt)
    → recherche/sélection préservée, fidélité visuelle 100%

Usage idéal :
  - PDFs Adobe InDesign/Photoshop avec layouts complexes
  - PDFs scannés où l'OCR n'est pas requis (le texte natif suffit)
  - Tout PDF design_tool où le rendu fluide casse la mise en page
"""

from __future__ import annotations

import logging
import os
import re
import tempfile
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)

_CONTROL_CHARS = re.compile(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]')


class HybridConverter:
    """
    Convertisseur "image + texte invisible" pour PDFs au layout complexe.

    Aucun LLM. Aucune reconstruction de flux : l'apparence est strictement
    identique à l'original.
    """

    def __init__(self, dpi: int = 300, text_overlay: bool = True) -> None:
        """
        Args:
            dpi          : résolution rendu pages (300 = qualité impression)
            text_overlay : superposer une couche texte invisible (recherche/copie)
        """
        self.dpi          = dpi
        self.text_overlay = text_overlay

    def convert(
        self,
        input_pdf:    str | Path,
        output_docx:  str | Path,
        pages:        list[int] | None = None,
    ) -> str:
        """
        Convertir un PDF en DOCX en préservant strictement l'apparence.

        Input  : chemin PDF + chemin DOCX de sortie
        Output : chemin DOCX généré
        """
        input_pdf   = str(input_pdf)
        output_docx = str(output_docx)
        logger.info("HybridConverter (dpi=%d) : %s -> %s",
                    self.dpi, input_pdf, output_docx)

        pdf_doc = fitz.open(input_pdf)
        doc     = Document()

        # Supprimer le paragraphe vide par défaut
        if doc.paragraphs:
            elem = doc.paragraphs[0]._element
            elem.getparent().remove(elem)

        page_numbers = pages if pages else list(range(len(pdf_doc)))

        with tempfile.TemporaryDirectory() as tmp_dir:
            for page_idx, page_num in enumerate(page_numbers):
                if page_num >= len(pdf_doc):
                    continue
                page = pdf_doc[page_num]
                logger.info("Page %d/%d", page_idx + 1, len(page_numbers))

                img_path = os.path.join(tmp_dir, f"page_{page_num}.png")
                self._render_page(page, img_path)

                w_in = page.rect.width  / 72.0
                h_in = page.rect.height / 72.0

                section = doc.sections[0] if page_idx == 0 else doc.add_section()
                section.page_width    = Inches(w_in)
                section.page_height   = Inches(h_in)
                section.left_margin   = Cm(0)
                section.right_margin  = Cm(0)
                section.top_margin    = Cm(0)
                section.bottom_margin = Cm(0)
                section.orientation   = (WD_ORIENT.LANDSCAPE if w_in > h_in
                                         else WD_ORIENT.PORTRAIT)

                para = doc.add_paragraph()
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after  = Pt(0)
                para.paragraph_format.line_spacing = 1.0
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.add_run().add_picture(img_path, width=Inches(w_in))

                if self.text_overlay:
                    self._add_text_overlay(doc, page)

        pdf_doc.close()
        doc.save(output_docx)
        logger.info("HybridConverter terminé : %s", output_docx)
        return output_docx

    # ── helpers ──────────────────────────────────────────────────────────────

    def _render_page(self, page: fitz.Page, output_path: str) -> None:
        zoom = self.dpi / 72.0
        pix  = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
        pix.save(output_path)

    def _add_text_overlay(self, doc: Document, page: fitz.Page) -> None:
        text_dict = page.get_text("dict", flags=fitz.TEXT_PRESERVE_WHITESPACE)

        lines: list[str] = []
        for block in text_dict.get("blocks", []):
            if block.get("type") != 0:
                continue
            for line in block.get("lines", []):
                line_text = "".join(span.get("text", "") for span in line.get("spans", []))
                cleaned   = self._sanitize(line_text.strip())
                if cleaned:
                    lines.append(cleaned)

        if not lines:
            return

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after  = Pt(0)
        run = para.add_run(self._sanitize(" ".join(lines)))
        run.font.size       = Pt(1)
        run.font.color.rgb  = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.name       = "Arial"

    @staticmethod
    def _sanitize(text: str) -> str:
        return _CONTROL_CHARS.sub("", text)
