"""
DocxEnhancer — Post-traitement DOCX après conversion PDF.

Porté depuis https://github.com/CHRISTMardochee/pdf2word
et généralisé pour tout contexte métier (suppression des couleurs spécifiques).

11 étapes d'amélioration :
  1. Correction des marges de section (right_margin=0 fréquent)
  2. Extraction des zones de texte flottantes en paragraphes normaux
  3. Fusion des paragraphes fragmentés
  4. Normalisation des espacements (cap à 24pt)
  5. Normalisation des indentations
  6. Dépliage des tableaux à cellule unique
  7. Suppression des paragraphes vides
  8. Normalisation des polices
  9. Correction de l'orientation des images (via matrices PDF)
 10. Espacement autour des images
 11. Restauration des bordures de tableaux
"""

from __future__ import annotations

import logging
import re
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from lxml import etree

logger = logging.getLogger(__name__)

_W_NS  = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_MC_NS = "{http://schemas.openxmlformats.org/markup-compatibility/2006}"
_WPS_NS = "{http://schemas.microsoft.com/office/word/2010/wordprocessingShape}"
_V_NS   = "{urn:schemas-microsoft-com:vml}"
_A_NS   = "{http://schemas.openxmlformats.org/drawingml/2006/main}"


class DocxEnhancer:
    """
    Post-traitement d'un DOCX converti depuis PDF.
    Améliore la qualité sans LLM — uniquement manipulation XML python-docx.
    """

    def enhance(
        self,
        docx_path: str | Path,
        output_path: str | Path | None = None,
        *,
        source_pdf_path: str | Path | None = None,
    ) -> str:
        """
        Améliorer un DOCX après conversion.

        Input  : chemin DOCX (+ optionnellement le PDF source pour l'orientation des images)
        Output : chemin DOCX amélioré
        """
        if output_path is None:
            output_path = docx_path
        docx_path   = str(docx_path)
        output_path = str(output_path)

        logger.info("DocxEnhancer : %s", docx_path)
        doc = Document(docx_path)

        self._fix_section_margins(doc)
        self._extract_textboxes(doc)
        self._merge_fragmented_paragraphs(doc)
        self._normalize_spacing(doc)
        self._normalize_indentation(doc)
        self._unwrap_single_cell_tables(doc)
        self._remove_empty_paragraphs(doc)
        self._normalize_fonts(doc)

        if source_pdf_path:
            self._fix_image_orientation(doc, str(source_pdf_path))

        self._fix_image_spacing(doc)
        self._restore_table_borders(doc)

        doc.save(output_path)
        logger.info("DocxEnhancer terminé : %s", output_path)
        return output_path

    # ── Étape 1 : Marges ─────────────────────────────────────────────────────

    def _fix_section_margins(self, doc: Document) -> None:
        default  = Cm(1.5)
        min_mgn  = Cm(0.5)
        for section in doc.sections:
            if section.right_margin is not None and section.right_margin < min_mgn:
                section.right_margin = default
            if section.left_margin is not None and section.left_margin < min_mgn:
                section.left_margin = default

    # ── Étape 2 : Zones de texte flottantes ──────────────────────────────────

    def _extract_textboxes(self, doc: Document) -> None:
        body     = doc.element.body
        to_extract: list[dict] = []

        for para in body.findall(qn("w:p")):
            drawings = para.findall(f".//{_MC_NS}AlternateContent")
            drawings += para.findall(".//" + qn("w:drawing"))

            for drawing in drawings:
                for txbx in drawing.findall(f".//{_WPS_NS}txbx"):
                    inner = txbx.find(f"{_W_NS}txbxContent")
                    if inner is not None:
                        inner_paras = inner.findall(f"{_W_NS}p")
                        if inner_paras:
                            to_extract.append({"parent": para, "inner": inner_paras})

            for shape in para.findall(f".//{_V_NS}shape"):
                txbx = shape.find(f"{_V_NS}textbox")
                if txbx is not None:
                    inner = txbx.find(f"{_W_NS}txbxContent")
                    if inner is not None:
                        inner_paras = inner.findall(f"{_W_NS}p")
                        if inner_paras:
                            to_extract.append({"parent": para, "inner": inner_paras})

        for item in to_extract:
            parent = item["parent"]
            try:
                idx = list(body).index(parent)
            except ValueError:
                continue
            for i, inner_para in enumerate(item["inner"]):
                body.insert(idx + 1 + i, deepcopy(inner_para))
            body.remove(parent)

        if to_extract:
            logger.info("Textboxes extraites : %d", len(to_extract))

    # ── Étape 3 : Fusion paragraphes fragmentés ───────────────────────────────

    def _merge_fragmented_paragraphs(self, doc: Document) -> None:
        body   = doc.element.body
        paras  = list(body.findall(qn("w:p")))
        merged = 0
        i = 0

        while i < len(paras) - 1:
            cur  = paras[i]
            nxt  = paras[i + 1]
            t1   = self._para_text(cur).strip()
            t2   = self._para_text(nxt).strip()

            if not t1 or not t2:
                i += 1
                continue

            if self._should_merge(cur, nxt, t1, t2):
                self._append_runs(cur, nxt)
                body.remove(nxt)
                paras.pop(i + 1)
                merged += 1
            else:
                i += 1

        if merged:
            logger.info("Paragraphes fusionnés : %d", merged)

    def _should_merge(self, p1, p2, t1: str, t2: str) -> bool:
        if t1 and t1[-1] in ".!?:;":
            return False
        if t2 and (t2[0] in "•·-–—" or re.match(r"^\d+[\.\)]\s", t2)):
            return False
        if t1 and t1[-1] in ",;":
            return True
        f1 = self._dominant_font(p1)
        f2 = self._dominant_font(p2)
        if f1 and f2 and f1 == f2 and t1 and t1[-1] not in ".!?:;":
            return len(t1) < 80
        return False

    def _para_text(self, para_el) -> str:
        return "".join(t.text or "" for t in para_el.findall(".//" + qn("w:t")))

    def _dominant_font(self, para_el) -> dict | None:
        runs = para_el.findall(qn("w:r"))
        if not runs:
            return None
        rpr = runs[0].find(qn("w:rPr"))
        if rpr is None:
            return {}
        info: dict = {}
        sz = rpr.find(qn("w:sz"))
        if sz is not None:
            info["size"] = sz.get(qn("w:val"))
        info["bold"]   = rpr.find(qn("w:b"))   is not None
        info["italic"] = rpr.find(qn("w:i"))   is not None
        rf = rpr.find(qn("w:rFonts"))
        if rf is not None:
            info["font"] = rf.get(qn("w:ascii")) or rf.get(qn("w:hAnsi"))
        return info

    def _append_runs(self, target, source) -> None:
        sp   = etree.SubElement(target, qn("w:r"))
        sp_t = etree.SubElement(sp, qn("w:t"))
        sp_t.text = " "
        sp_t.set(qn("xml:space"), "preserve")
        for run in source.findall(qn("w:r")):
            target.append(deepcopy(run))

    # ── Étape 4 : Normalisation espacements ───────────────────────────────────

    def _normalize_spacing(self, doc: Document) -> None:
        cap = Pt(24)
        for para in doc.paragraphs:
            pf = para.paragraph_format
            if pf.space_before is not None and pf.space_before > cap:
                pf.space_before = cap
            if pf.space_after is not None and pf.space_after > cap:
                pf.space_after = cap

    # ── Étape 5 : Normalisation indentations ──────────────────────────────────

    def _normalize_indentation(self, doc: Document) -> None:
        noise = Cm(0.3)
        cap   = Cm(3.0)
        for para in doc.paragraphs:
            pf = para.paragraph_format
            if pf.first_line_indent is not None and pf.first_line_indent < 0:
                continue
            if pf.left_indent is not None:
                if 0 < pf.left_indent < noise:
                    pf.left_indent = Pt(0)
                elif pf.left_indent > cap and not para.text.strip().startswith(("•", "-", "–")):
                    pf.left_indent = cap
            if pf.right_indent is not None and 0 < pf.right_indent < noise:
                pf.right_indent = Pt(0)

    # ── Étape 6 : Tableaux à cellule unique ───────────────────────────────────

    def _unwrap_single_cell_tables(self, doc: Document) -> None:
        body = doc.element.body
        for table in doc.tables:
            if not (len(table.rows) == 1 and len(table.columns) == 1):
                continue
            tbl = table._tbl
            try:
                idx = list(body).index(tbl)
            except ValueError:
                continue
            for i, cell_para in enumerate(table.cell(0, 0).paragraphs):
                body.insert(idx + 1 + i, deepcopy(cell_para._element))
            body.remove(tbl)

    # ── Étape 7 : Paragraphes vides ───────────────────────────────────────────

    def _remove_empty_paragraphs(self, doc: Document) -> None:
        body = doc.element.body
        for para in list(body.findall(qn("w:p"))):
            text      = self._para_text(para).strip()
            has_media = bool(
                para.findall(".//" + qn("w:drawing"))
                or para.findall(".//" + qn("w:pict"))
            )
            if text or has_media:
                continue
            # Conserver les sauts de page
            br = para.findall(".//" + qn("w:br"))
            if any(b.get(qn("w:type")) == "page" for b in br):
                continue
            ppr = para.find(qn("w:pPr"))
            if ppr is not None and ppr.find(qn("w:sectPr")) is not None:
                continue
            body.remove(para)

    # ── Étape 8 : Normalisation polices ───────────────────────────────────────

    def _normalize_fonts(self, doc: Document) -> None:
        size_counts: dict = {}
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size:
                    size_counts[run.font.size] = size_counts.get(run.font.size, 0) + len(run.text)

        if not size_counts:
            return
        dominant = max(size_counts, key=size_counts.get)
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.size and abs(run.font.size - dominant) <= Pt(0.5):
                    run.font.size = dominant

    # ── Étape 9 : Orientation images ─────────────────────────────────────────

    def _fix_image_orientation(self, doc: Document, source_pdf_path: str) -> None:
        import fitz
        pdf_doc    = fitz.open(source_pdf_path)
        transforms = []
        try:
            for page in pdf_doc:
                for info in page.get_image_info(xrefs=True):
                    t = info.get("transform")
                    if t and len(t) >= 4:
                        flip_h = t[0] < 0
                        flip_v = t[3] > 0
                        if flip_h or flip_v:
                            transforms.append({"flip_h": flip_h, "flip_v": flip_v})
        finally:
            pdf_doc.close()

        if not transforms:
            return

        fixed = 0
        for para in doc.paragraphs:
            for run in para.runs:
                for drawing in run._element.findall(".//" + qn("w:drawing")):
                    for xfrm in drawing.findall(f".//{_A_NS}xfrm"):
                        if fixed < len(transforms):
                            t = transforms[fixed]
                            if t["flip_h"]:
                                xfrm.set("flipH", "1")
                            if t["flip_v"]:
                                xfrm.set("flipV", "1")
                            fixed += 1
        if fixed:
            logger.info("Orientation corrigée pour %d images", fixed)

    # ── Étape 10 : Espacement images ─────────────────────────────────────────

    def _fix_image_spacing(self, doc: Document) -> None:
        spacing = Pt(12)
        for para in doc.paragraphs:
            has_img = bool(
                para._element.findall(".//" + qn("w:drawing"))
                or para._element.findall(".//" + qn("w:pict"))
            )
            if has_img:
                pf = para.paragraph_format
                if pf.space_before is None or pf.space_before < spacing:
                    pf.space_before = spacing
                if pf.space_after is None or pf.space_after < spacing:
                    pf.space_after = spacing

    # ── Étape 11 : Bordures tableaux ─────────────────────────────────────────

    def _restore_table_borders(self, doc: Document) -> None:
        """
        Restaurer les bordures sur tous les tableaux.
        Couleur adaptée : gris clair pour tableaux de données.
        """
        BORDER_COLOR = "D0D0D0"
        BORDER_SIZE  = "4"

        for table in doc.tables:
            has_content = any(
                p.text.strip()
                for row in table.rows
                for cell in row.cells
                for p in cell.paragraphs
            )
            if not has_content:
                continue

            tbl    = table._tbl
            tbl_pr = tbl.find(f"{_W_NS}tblPr")
            if tbl_pr is None:
                tbl_pr = etree.SubElement(tbl, f"{_W_NS}tblPr")

            existing = tbl_pr.find(f"{_W_NS}tblBorders")
            if existing is not None:
                tbl_pr.remove(existing)

            borders = etree.SubElement(tbl_pr, f"{_W_NS}tblBorders")
            for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                border = etree.SubElement(borders, f"{_W_NS}{edge}")
                border.set(f"{_W_NS}val",   "single")
                border.set(f"{_W_NS}sz",    BORDER_SIZE)
                border.set(f"{_W_NS}space", "0")
                border.set(f"{_W_NS}color", BORDER_COLOR)
