"""
DoclingConverter — Conversion PDF vers DOCX via IBM Docling avec extraction
maximale des éléments visuels (images, tableaux, structure).

100% open-source, 100% offline (après téléchargement des modèles ML),
0 dépendance cloud, aucune clé API.

Pipeline complet activé :
  - generate_picture_images = True   → extrait toutes les images embarquées
  - generate_page_images    = True   → rendu HD de chaque page (référence visuelle)
  - do_table_structure      = True   → reconnaissance structure tableaux (TableFormer)
  - do_ocr                  = True   → OCR pour zones scannées
  - do_picture_classification = True → classifie les images (logo, photo, schéma…)
  - images_scale            = 2.0    → résolution 2× (qualité impression)

L'output DOCX intègre :
  - Hiérarchie des titres (Heading 1/2/3) détectée par ML
  - Listes à puces et numérotées
  - Tableaux natifs Word (cellules, headers)
  - Images insérées aux bonnes positions dans le flux
  - Polices, tailles, couleurs et fonds extraits du PDF source via PyMuPDF
"""

from __future__ import annotations

import io
import logging
import os
import tempfile
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

logger = logging.getLogger(__name__)


class DoclingConverter:
    """
    Convertisseur PDF → DOCX via IBM Docling avec pipeline visuel complet.

    100% offline. Aucun appel cloud. Aucun LLM génératif (modèles ML
    spécialisés pour la compréhension document : layout, tables, OCR).
    """

    def __init__(
        self,
        *,
        images_scale:        float = 2.0,
        do_ocr:              bool  = True,
        do_table_structure:  bool  = True,
        extract_pictures:    bool  = True,
        classify_pictures:   bool  = False,  # nécessite MSVC sur Windows (torch.compile)
    ) -> None:
        try:
            from docling.document_converter import DocumentConverter, PdfFormatOption
            from docling.datamodel.base_models import InputFormat
            from docling.datamodel.pipeline_options import PdfPipelineOptions
        except ImportError as exc:
            raise ImportError(
                "DoclingConverter nécessite IBM Docling. Installer :\n"
                "    pip install docling\n"
                "ou : pip install docpipeline[docling]"
            ) from exc

        # Activer toutes les fonctionnalités visuelles avancées
        opts = PdfPipelineOptions()
        opts.images_scale              = images_scale
        opts.generate_page_images      = True
        opts.generate_picture_images   = extract_pictures
        opts.do_table_structure        = do_table_structure
        opts.do_ocr                    = do_ocr
        opts.do_picture_classification = classify_pictures
        opts.generate_table_images     = True

        self._converter = DocumentConverter(
            format_options={
                InputFormat.PDF: PdfFormatOption(pipeline_options=opts)
            }
        )

    def convert(
        self,
        input_pdf:    str | Path,
        output_docx:  str | Path,
        pages:        list[int] | None = None,
    ) -> str:
        """
        Convertir un PDF en DOCX avec extraction visuelle maximale.

        Input  : chemin PDF + chemin DOCX
        Output : chemin DOCX généré
        """
        import fitz

        input_pdf  = os.path.abspath(str(input_pdf))
        output_pdf = os.path.abspath(str(output_docx))

        if not os.path.isfile(input_pdf):
            raise FileNotFoundError(f"PDF introuvable : {input_pdf}")

        logger.info("DoclingConverter (full pipeline) : %s -> %s", input_pdf, output_pdf)

        actual_pdf = input_pdf
        tmp_pdf    = None
        if pages is not None:
            tmp_pdf = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
            tmp_pdf.close()
            src = fitz.open(input_pdf)
            dst = fitz.open()
            try:
                for p in pages:
                    if p < len(src):
                        dst.insert_pdf(src, from_page=p, to_page=p)
                dst.save(tmp_pdf.name)
            finally:
                dst.close()
                src.close()
            actual_pdf = tmp_pdf.name

        try:
            logger.info("Analyse Docling ML (modèles téléchargés au 1er run)...")
            result      = self._converter.convert(actual_pdf)
            docling_doc = result.document

            doc = self._build_docx_from_doc(docling_doc)

            logger.info("Application des styles visuels du PDF source...")
            self._apply_pdf_visual_styles(doc, actual_pdf)

            for section in doc.sections:
                section.left_margin   = Cm(2.0)
                section.right_margin  = Cm(2.0)
                section.top_margin    = Cm(1.5)
                section.bottom_margin = Cm(1.5)

            doc.save(output_pdf)
            logger.info("DoclingConverter terminé : %s", output_pdf)
            return output_pdf

        finally:
            if tmp_pdf is not None:
                try:
                    os.unlink(tmp_pdf.name)
                except OSError:
                    pass

    # ── Construction DOCX en itérant sur le DoclingDocument ──────────────────

    def _build_docx_from_doc(self, docling_doc: Any) -> Document:
        """
        Itère sur les éléments structurés (TextItem, PictureItem, TableItem)
        du DoclingDocument et construit un DOCX préservant l'ordre et la
        hiérarchie sémantique détectée par les modèles ML.
        """
        from docling_core.types.doc import (
            PictureItem, SectionHeaderItem, TableItem, TextItem, ListItem
        )

        doc = Document()
        style              = doc.styles["Normal"]
        style.font.name    = "Calibri"
        style.font.size    = Pt(10)

        n_text  = 0
        n_imgs  = 0
        n_tabs  = 0
        n_heads = 0

        for item, level in docling_doc.iterate_items():
            # ── Titres / Section headers ─────────────────────────────────────
            if isinstance(item, SectionHeaderItem):
                heading_level = min(max(item.level or 1, 1), 6)
                doc.add_heading(item.text, level=heading_level)
                n_heads += 1
                continue

            # ── Listes ───────────────────────────────────────────────────────
            if isinstance(item, ListItem):
                style_name = "List Bullet"
                if hasattr(item, "marker") and item.marker and item.marker.strip().rstrip(".)").isdigit():
                    style_name = "List Number"
                para = doc.add_paragraph(item.text, style=style_name)
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after  = Pt(1)
                n_text += 1
                continue

            # ── Texte / paragraphe ───────────────────────────────────────────
            if isinstance(item, TextItem):
                text = item.text.strip()
                if not text:
                    continue
                # Détection titre par convention label
                label = getattr(item, "label", None)
                label_str = str(label).lower() if label else ""

                if "title" in label_str:
                    doc.add_heading(text, level=1)
                    n_heads += 1
                elif "section_header" in label_str or "header" in label_str:
                    doc.add_heading(text, level=2)
                    n_heads += 1
                elif "list_item" in label_str:
                    para = doc.add_paragraph(text, style="List Bullet")
                    para.paragraph_format.space_before = Pt(1)
                    para.paragraph_format.space_after  = Pt(1)
                    n_text += 1
                elif "caption" in label_str:
                    para = doc.add_paragraph(text)
                    for run in para.runs:
                        run.font.italic = True
                        run.font.size   = Pt(9)
                    n_text += 1
                else:
                    para = doc.add_paragraph(text)
                    para.paragraph_format.space_before = Pt(2)
                    para.paragraph_format.space_after  = Pt(2)
                    n_text += 1
                continue

            # ── Image extraite ───────────────────────────────────────────────
            if isinstance(item, PictureItem):
                self._insert_picture(doc, item, docling_doc)
                n_imgs += 1
                continue

            # ── Tableau natif ────────────────────────────────────────────────
            if isinstance(item, TableItem):
                self._insert_table(doc, item)
                n_tabs += 1
                continue

        logger.info("DOCX construit : %d titres, %d paragraphes, %d images, %d tableaux",
                    n_heads, n_text, n_imgs, n_tabs)
        return doc

    def _insert_picture(self, doc: Document, item: Any, docling_doc: Any) -> None:
        """Insère une image extraite par Docling dans le DOCX."""
        try:
            pil_image = item.get_image(docling_doc) if hasattr(item, "get_image") else None
            if pil_image is None and hasattr(item, "image") and item.image:
                pil_image = item.image.pil_image if hasattr(item.image, "pil_image") else None

            if pil_image is None:
                # Fallback : caption seulement
                if hasattr(item, "caption_text") and item.caption_text:
                    para = doc.add_paragraph()
                    run = para.add_run(f"[Image : {item.caption_text(docling_doc)}]")
                    run.font.italic = True
                    run.font.size   = Pt(9)
                return

            # Sauvegarder l'image en mémoire et l'insérer
            buf = io.BytesIO()
            pil_image.save(buf, format="PNG")
            buf.seek(0)

            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run()

            # Largeur adaptée à la page (max 6 pouces)
            max_width_in = 6.0
            img_w_px     = pil_image.width
            target_in    = min(max_width_in, img_w_px / 100.0)
            run.add_picture(buf, width=Inches(target_in))

            # Caption si présente
            if hasattr(item, "caption_text"):
                try:
                    cap = item.caption_text(docling_doc)
                    if cap:
                        cap_para = doc.add_paragraph()
                        cap_run  = cap_para.add_run(cap)
                        cap_run.font.italic = True
                        cap_run.font.size   = Pt(9)
                        cap_para.alignment  = WD_ALIGN_PARAGRAPH.CENTER
                except Exception:
                    pass

            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after  = Pt(6)
        except Exception as exc:
            logger.warning("Insertion image échouée : %s", exc)

    def _insert_table(self, doc: Document, item: Any) -> None:
        """Insère un tableau Docling comme tableau natif Word."""
        try:
            # Récupérer les données (différentes API selon version)
            if hasattr(item, "export_to_dataframe"):
                df = item.export_to_dataframe()
                rows_data = [df.columns.tolist()] + df.values.tolist()
            elif hasattr(item, "data") and item.data:
                grid = item.data.grid if hasattr(item.data, "grid") else item.data
                rows_data = [[cell.text if hasattr(cell, "text") else str(cell)
                              for cell in row] for row in grid]
            else:
                return

            if not rows_data or not rows_data[0]:
                return

            row_count = len(rows_data)
            col_count = max(len(r) for r in rows_data)

            table = doc.add_table(rows=row_count, cols=col_count)
            table.autofit = True
            self._set_table_borders(table)

            for r, row in enumerate(rows_data):
                for c in range(min(len(row), col_count)):
                    cell      = table.cell(r, c)
                    cell_text = str(row[c]) if row[c] is not None else ""
                    para      = cell.paragraphs[0]
                    if cell_text:
                        run = para.add_run(cell_text)
                        run.font.size = Pt(9)
                        run.font.name = "Calibri"
                        if r == 0:
                            run.font.bold       = True
                            run.font.color.rgb  = RGBColor(255, 255, 255)

                if r == 0:
                    for c in range(min(len(row), col_count)):
                        cell = table.cell(0, c)
                        shd  = OxmlElement("w:shd")
                        shd.set(qn("w:fill"), "2E75B6")
                        shd.set(qn("w:val"),  "clear")
                        cell._tc.get_or_add_tcPr().append(shd)
        except Exception as exc:
            logger.warning("Insertion tableau échouée : %s", exc)

    # ── Post-traitement visuel via PyMuPDF (couleurs, fonds, polices) ────────

    def _apply_pdf_visual_styles(self, doc: Document, pdf_path: str) -> None:
        import fitz
        pdf_doc = fitz.open(pdf_path)
        try:
            visual = self._extract_visual_data(pdf_doc)
        finally:
            pdf_doc.close()

        self._apply_fonts_to_paragraphs(doc, visual)
        self._apply_heading_backgrounds(doc, visual)

    def _extract_visual_data(self, pdf_doc: Any) -> dict:
        import fitz
        data: dict = {"text_blocks": [], "backgrounds": [], "font_stats": {}}
        size_counts: dict[float, int] = {}

        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            for block in page.get_text("dict").get("blocks", []):
                if block.get("type") != 0:
                    continue
                for line in block.get("lines", []):
                    for span in line.get("spans", []):
                        text = span.get("text", "").strip()
                        if not text:
                            continue
                        size      = round(span.get("size", 10), 1)
                        flags     = span.get("flags", 0)
                        color_int = span.get("color", 0)
                        r = (color_int >> 16) & 0xFF
                        g = (color_int >>  8) & 0xFF
                        b = color_int & 0xFF
                        data["text_blocks"].append({
                            "page":   page_num,
                            "text":   text[:80],
                            "size":   size,
                            "bold":   bool(flags & (1 << 4)),
                            "italic": bool(flags & (1 << 1)),
                            "color":  f"{r:02X}{g:02X}{b:02X}",
                            "bbox":   span.get("bbox", []),
                        })
                        size_counts[size] = size_counts.get(size, 0) + len(text)

            for d in page.get_drawings():
                fill = d.get("fill")
                rect = d.get("rect")
                if fill is None or rect is None:
                    continue
                dr = fitz.Rect(rect)
                if dr.width < 50 or dr.height < 10:
                    continue
                hex_c = self._rgb_to_hex(fill)
                if hex_c == "FFFFFF":
                    continue
                data["backgrounds"].append({
                    "page": page_num, "rect": rect,
                    "width": dr.width, "height": dr.height,
                    "color": hex_c, "y": dr.y0,
                })

        data["font_stats"]["body_size"] = (
            max(size_counts, key=size_counts.get) if size_counts else 10.0
        )
        return data

    def _apply_fonts_to_paragraphs(self, doc: Document, visual: dict) -> None:
        body_size  = visual["font_stats"].get("body_size", 10.0)
        text_fonts: dict[str, dict] = {}
        for block in visual["text_blocks"]:
            key = block["text"][:40].strip().lower()
            if key and key not in text_fonts:
                text_fonts[key] = block

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            match = text_fonts.get(text[:40].strip().lower())
            if not match:
                continue
            for run in para.runs:
                pdf_size = match["size"]
                run.font.size = Pt(pdf_size if pdf_size > body_size * 1.3 else body_size)
                if match["bold"]:
                    run.font.bold = True
                if match["italic"]:
                    run.font.italic = True
                color = match["color"]
                if color != "000000":
                    run.font.color.rgb = RGBColor(
                        int(color[0:2], 16), int(color[2:4], 16), int(color[4:6], 16)
                    )

    def _apply_heading_backgrounds(self, doc: Document, visual: dict) -> None:
        backgrounds = visual["backgrounds"]
        if not backgrounds:
            return
        header_bands = [bg for bg in backgrounds
                        if bg["width"] > 200 and 15 < bg["height"] < 50]
        info_boxes   = [bg for bg in backgrounds
                        if bg["width"] > 200 and bg["height"] > 50]

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            for block in visual["text_blocks"]:
                btext = block["text"].strip()
                if not btext or text[:30].lower() != btext[:30].lower():
                    continue
                bbox = block.get("bbox", [])
                if len(bbox) < 4:
                    continue
                ty, tp = bbox[1], block["page"]
                applied = False
                for band in header_bands:
                    if band["page"] != tp:
                        continue
                    br = band["rect"]
                    if (br.y0 - 5) <= ty <= (br.y1 + 5):
                        self._set_paragraph_shading(para, band["color"])
                        if self._is_dark_color(band["color"]):
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(255, 255, 255)
                                run.font.bold      = True
                        applied = True
                        break
                if not applied:
                    for box in info_boxes:
                        if box["page"] != tp:
                            continue
                        br = box["rect"]
                        if (br.y0 - 5) <= ty <= (br.y1 + 5):
                            self._set_paragraph_shading(para, box["color"])
                            break
                break

    # ── Helpers ──────────────────────────────────────────────────────────────

    @staticmethod
    def _set_paragraph_shading(para: Any, hex_color: str) -> None:
        pPr = para._element.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), hex_color)
        shd.set(qn("w:val"),  "clear")
        pPr.append(shd)

    @staticmethod
    def _is_dark_color(hex_color: str) -> bool:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return (0.299 * r + 0.587 * g + 0.114 * b) / 255 < 0.5

    @staticmethod
    def _rgb_to_hex(rgb: tuple) -> str:
        r = int(min(max(rgb[0], 0), 1) * 255)
        g = int(min(max(rgb[1], 0), 1) * 255)
        b = int(min(max(rgb[2], 0), 1) * 255)
        return f"{r:02X}{g:02X}{b:02X}"

    def _set_table_borders(self, table: Any) -> None:
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        tbl  = table._tbl
        tbl_pr = tbl.find(f"{{{W_NS}}}tblPr")
        if tbl_pr is None:
            tbl_pr = OxmlElement("w:tblPr")
            tbl.insert(0, tbl_pr)
        existing = tbl_pr.find(f"{{{W_NS}}}tblBorders")
        if existing is not None:
            tbl_pr.remove(existing)
        borders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            border = OxmlElement(f"w:{edge}")
            border.set(qn("w:val"),   "single")
            border.set(qn("w:sz"),    "4")
            border.set(qn("w:color"), "D0D0D0")
            border.set(qn("w:space"), "0")
            borders.append(border)
        tbl_pr.append(borders)
