"""
parse_word.py — Parsing complet d'un .docx avec préservation des styles.

Symétrique de `parse_pdf.py` côté Word. Une seule ouverture python-docx, sortie
unifiée en DataFrames + doc_summary.

Pattern transversal du projet (cf. Kezhan, call 2026-05-08) :

    extract  : parse_word(docx)             → spans avec styles + span_id stable
    modify   : on change SEULEMENT le text de chaque span (translate, redact, …)
    rebuild  : apply_changes(docx_in, {span_id: new_text}, docx_out)
               → reconstruit le .docx avec les nouveaux textes mais styles intacts

Le `span_id` est la **clé stable** qui fait le pont entre extraction et
reconstruction. Format : `w_<para_idx>_<run_idx>` — déterministe, reproductible,
résiste aux modifications de texte.

Sortie de `parse_word(docx)` :
    {
        "paragraph_df": DataFrame,   # 1 ligne = 1 paragraphe (texte, style, alignement, indent, ...)
        "span_df":      DataFrame,   # 1 ligne = 1 run (span_id, text, font, bold, italic, color, ...)
        "image_df":     DataFrame,   # 1 ligne = 1 image embarquée (taille, position)
        "table_df":     DataFrame,   # 1 ligne = 1 table (n_rows, n_cols, styles)
        "doc_summary":  dict,        # JSON synthèse : metadata, comptes, signaux structurels
    }

Aucun LLM (règle CLAUDE.md : LLM réservé à translation / summarization /
Excel SQL agent ; le parsing est intégralement heuristique).
"""

from __future__ import annotations

import hashlib
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Optional

import pandas as pd
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║ 1. CONSTANTES                                                              ║
# ╚════════════════════════════════════════════════════════════════════════════╝

_ALIGNMENT_NAMES = {
    0: "left", 1: "center", 2: "right", 3: "justify",
    4: "distribute", 5: "justify_med", 7: "justify_hi", 8: "justify_low",
    9: "thai_justify",
}

_LINE_SPACING_RULES = {
    1: "single", 2: "1.5", 3: "double", 4: "at_least", 5: "exactly", 6: "multiple",
}

_UNDERLINE_NAMES = {
    True: "single", False: "none",
    1: "single", 2: "words", 3: "double", 4: "dotted", 5: "thick", 6: "dash",
    7: "dot_dash", 8: "dot_dot_dash", 9: "wavy",
}


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║ 2. HELPERS — petites lectures sûres                                        ║
# ╚════════════════════════════════════════════════════════════════════════════╝

def _sha256(path: Path) -> str:
    h = hashlib.sha256()
    with open(path, "rb") as f:
        for chunk in iter(lambda: f.read(1 << 20), b""):
            h.update(chunk)
    return h.hexdigest()


def _safe(getter, default=None):
    """Appel getter ; retourne default si exception ou None."""
    try:
        v = getter()
        return v if v is not None else default
    except Exception:
        return default


def _alignment_name(align) -> Optional[str]:
    if align is None:
        return None
    try:
        return _ALIGNMENT_NAMES.get(int(align), str(align))
    except Exception:
        return str(align)


def _underline_name(u) -> Optional[str]:
    if u is None:
        return None
    return _UNDERLINE_NAMES.get(u, str(u))


def _color_to_hex(color) -> Optional[str]:
    """Conversion robuste : RGB direct, ou theme color, ou None."""
    if color is None:
        return None
    try:
        if color.rgb is not None:
            return f"#{color.rgb}"
    except Exception:
        pass
    try:
        if color.theme_color is not None:
            return f"theme:{color.theme_color}"
    except Exception:
        pass
    return None


def _highlight_name(h) -> Optional[str]:
    if h is None:
        return None
    try:
        return str(h).split('.')[-1].lower()
    except Exception:
        return None


def _emu_to_pt(emu) -> Optional[float]:
    """English Metric Units (1 pt = 12700 EMU)."""
    if emu is None:
        return None
    try:
        return round(int(emu) / 12700.0, 2)
    except Exception:
        return None


def _twips_to_pt(twips) -> Optional[float]:
    """Twentieths of a point (1 pt = 20 twips)."""
    if twips is None:
        return None
    try:
        return round(int(twips) / 20.0, 2)
    except Exception:
        return None


def _hyperlink_target(run) -> Optional[str]:
    """Récupère la cible d'hyperlink si le run est dans un <w:hyperlink>."""
    try:
        parent = run._element.getparent()
        if parent is not None and parent.tag == qn("w:hyperlink"):
            r_id = parent.get(qn("r:id"))
            if r_id and run.part is not None:
                rels = run.part.rels
                if r_id in rels:
                    return rels[r_id].target_ref
    except Exception:
        pass
    return None


def _list_info(para: Paragraph) -> tuple[Optional[int], Optional[str]]:
    """Niveau de liste + référence numbering (None si pas une liste)."""
    try:
        pPr = para._p.find(qn("w:pPr"))
        if pPr is None:
            return None, None
        numPr = pPr.find(qn("w:numPr"))
        if numPr is None:
            return None, None
        ilvl = numPr.find(qn("w:ilvl"))
        numId = numPr.find(qn("w:numId"))
        level = int(ilvl.get(qn("w:val"))) if ilvl is not None else None
        num_id = numId.get(qn("w:val")) if numId is not None else None
        return level, num_id
    except Exception:
        return None, None


def _is_track_change_run(run) -> tuple[bool, bool]:
    """(is_insertion, is_deletion) — détecte si le run est dans un <w:ins> ou <w:del>."""
    try:
        parent = run._element.getparent()
        while parent is not None:
            if parent.tag == qn("w:ins"):
                return True, False
            if parent.tag == qn("w:del"):
                return False, True
            parent = parent.getparent()
    except Exception:
        pass
    return False, False


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║ 3. EXTRACTORS — paragraphes, spans, tables, images, sections               ║
# ╚════════════════════════════════════════════════════════════════════════════╝

def _extract_paragraph(para: Paragraph, para_idx: int) -> dict:
    """Extraction COMPLÈTE des propriétés d'un paragraphe."""
    pf = para.paragraph_format
    style = para.style if para.style else None
    list_level, list_num_id = _list_info(para)

    return {
        "paragraph_index":    para_idx,
        "text":               para.text,
        "char_count":         len(para.text or ""),
        "style_name":         style.name if style else "",
        "style_type":         _safe(lambda: str(style.type).split('.')[-1].lower(), "") if style else "",
        "heading_level":      _heading_level(para),
        "alignment":          _alignment_name(pf.alignment),
        "first_line_indent":  _safe(lambda: pf.first_line_indent.pt if pf.first_line_indent else None),
        "left_indent":        _safe(lambda: pf.left_indent.pt if pf.left_indent else None),
        "right_indent":       _safe(lambda: pf.right_indent.pt if pf.right_indent else None),
        "space_before":       _safe(lambda: pf.space_before.pt if pf.space_before else None),
        "space_after":        _safe(lambda: pf.space_after.pt if pf.space_after else None),
        "line_spacing":       _safe(lambda: float(pf.line_spacing) if pf.line_spacing else None),
        "line_spacing_rule":  _LINE_SPACING_RULES.get(_safe(lambda: int(pf.line_spacing_rule)) if pf.line_spacing_rule else None),
        "keep_together":      _safe(lambda: bool(pf.keep_together)),
        "keep_with_next":     _safe(lambda: bool(pf.keep_with_next)),
        "page_break_before":  _safe(lambda: bool(pf.page_break_before)),
        "widow_control":      _safe(lambda: bool(pf.widow_control)),
        "list_level":         list_level,
        "list_num_id":        list_num_id,
        "is_list_item":       list_level is not None,
        "n_runs":             len(para.runs),
    }


def _extract_span(
    run,
    para_idx: int,
    run_idx: int,
    *,
    span_id: Optional[str] = None,
    in_table: bool = False,
    table_pos: Optional[tuple[int, int, int]] = None,
) -> dict:
    """
    Extraction COMPLÈTE des propriétés d'un run (= span Word).

    `span_id` :
      - hors table : `w_<para>_<run>`
      - dans table : `w_t_<table>_<row>_<col>_<para>_<run>` (passé via span_id=)
    """
    font = run.font
    is_ins, is_del = _is_track_change_run(run)
    return {
        "span_id":          span_id or f"w_{para_idx}_{run_idx}",
        "in_table":         in_table,
        "table_index":      table_pos[0] if table_pos else None,
        "row_index":        table_pos[1] if table_pos else None,
        "col_index":        table_pos[2] if table_pos else None,
        "paragraph_index":  para_idx,
        "run_index":        run_idx,
        "text":             run.text or "",
        "char_count":       len(run.text or ""),
        # Font basics
        "font_name":        font.name or "",
        "font_size_pt":     _safe(lambda: font.size.pt if font.size else None),
        # Weight & emphasis
        "bold":             _safe(lambda: bool(run.bold)),
        "italic":           _safe(lambda: bool(run.italic)),
        "underline":        _underline_name(_safe(lambda: run.underline)),
        "strike":           _safe(lambda: bool(font.strike)),
        "double_strike":    _safe(lambda: bool(font.double_strike)),
        # Color
        "color":            _color_to_hex(font.color) if font.color is not None else None,
        "highlight":        _highlight_name(_safe(lambda: font.highlight_color)),
        # Position & casing
        "subscript":        _safe(lambda: bool(font.subscript)),
        "superscript":      _safe(lambda: bool(font.superscript)),
        "all_caps":         _safe(lambda: bool(font.all_caps)),
        "small_caps":       _safe(lambda: bool(font.small_caps)),
        # Effects
        "shadow":           _safe(lambda: bool(font.shadow)),
        "outline":          _safe(lambda: bool(font.outline)),
        "emboss":           _safe(lambda: bool(font.emboss)),
        "imprint":          _safe(lambda: bool(font.imprint)),
        # Style
        "char_style":       _safe(lambda: run.style.name if run.style else "", ""),
        # Hyperlink
        "hyperlink":        _hyperlink_target(run),
        # Track changes
        "is_insertion":     is_ins,
        "is_deletion":      is_del,
    }


def _extract_image(shape, image_idx: int) -> dict:
    """Métadonnées d'une image embarquée (inline shape)."""
    return {
        "image_index":  image_idx,
        "type":         _safe(lambda: str(shape.type).split('.')[-1].lower(), ""),
        "width_pt":     _safe(lambda: shape.width.pt if shape.width else None),
        "height_pt":    _safe(lambda: shape.height.pt if shape.height else None),
        "filename":     _safe(lambda: getattr(shape, "filename", "")),
    }


def _extract_table(table, table_idx: int) -> dict:
    """Métadonnées d'une table native + contenu cellules."""
    rows = list(table.rows)
    n_rows = len(rows)
    n_cols = len(rows[0].cells) if n_rows else 0
    cells = [[cell.text for cell in row.cells] for row in rows]
    return {
        "table_index":  table_idx,
        "n_rows":       n_rows,
        "n_cols":       n_cols,
        "style_name":   _safe(lambda: table.style.name if table.style else "", ""),
        "alignment":    _alignment_name(_safe(lambda: table.alignment)),
        "autofit":      _safe(lambda: bool(table.autofit)),
        "cells":        cells,                                # liste de listes (utile pour debug)
        "n_cells_with_text": sum(1 for row in cells for c in row if c.strip()),
    }


def _extract_section(section, section_idx: int) -> dict:
    """Section break = page setup + headers/footers."""
    return {
        "section_index":           section_idx,
        "page_width_pt":           _safe(lambda: section.page_width.pt if section.page_width else None),
        "page_height_pt":          _safe(lambda: section.page_height.pt if section.page_height else None),
        "orientation":             _safe(lambda: str(section.orientation).split('.')[-1].lower(), ""),
        "top_margin_pt":           _safe(lambda: section.top_margin.pt if section.top_margin else None),
        "bottom_margin_pt":        _safe(lambda: section.bottom_margin.pt if section.bottom_margin else None),
        "left_margin_pt":          _safe(lambda: section.left_margin.pt if section.left_margin else None),
        "right_margin_pt":         _safe(lambda: section.right_margin.pt if section.right_margin else None),
        "header_distance_pt":      _safe(lambda: section.header_distance.pt if section.header_distance else None),
        "footer_distance_pt":      _safe(lambda: section.footer_distance.pt if section.footer_distance else None),
        "different_first_page":    _safe(lambda: bool(section.different_first_page_header_footer)),
        "header_text":             _safe(lambda: " ".join(p.text for p in section.header.paragraphs).strip(), ""),
        "footer_text":             _safe(lambda: " ".join(p.text for p in section.footer.paragraphs).strip(), ""),
    }


def _heading_level(para: Paragraph) -> Optional[int]:
    """Niveau de heading (0 = Title, 1-9 = Heading 1-9, None sinon)."""
    style_name = para.style.name if para.style else ""
    s = (style_name or "").lower()
    if s == "title":
        return 0
    if s.startswith("heading "):
        try:
            return int(s.split()[-1])
        except Exception:
            return None
    return None


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║ 4. METADATA & SIGNAUX STRUCTURELS                                          ║
# ╚════════════════════════════════════════════════════════════════════════════╝

def _extract_doc_metadata(doc: Document) -> dict:
    """Properties core + custom du document."""
    cp = doc.core_properties
    return {
        "title":              cp.title or "",
        "author":             cp.author or "",
        "last_modified_by":   cp.last_modified_by or "",
        "subject":            cp.subject or "",
        "keywords":           cp.keywords or "",
        "category":           cp.category or "",
        "comments":           cp.comments or "",
        "revision":           cp.revision,
        "created":            cp.created.isoformat() if cp.created else None,
        "modified":           cp.modified.isoformat() if cp.modified else None,
        "last_printed":       cp.last_printed.isoformat() if cp.last_printed else None,
        "version":            cp.version or "",
        "content_status":     cp.content_status or "",
        "language":           cp.language or "",
    }


def _detect_source_tool(metadata: dict) -> str:
    """Outil ayant produit le .docx, depuis les metadata."""
    author = (metadata.get("author") or "").lower()
    last_mod = (metadata.get("last_modified_by") or "").lower()
    blob = " ".join([author, last_mod])
    if "libreoffice" in blob or "openoffice" in blob:
        return "LibreOffice"
    if "google" in blob:
        return "Google Docs"
    if "wps" in blob or "kingsoft" in blob:
        return "WPS Office"
    if "onlyoffice" in blob:
        return "OnlyOffice"
    if "pages" in blob:
        return "Apple Pages"
    # Microsoft Word ne se nomme pas explicitement dans les metadata —
    # on prend "Microsoft Word" par défaut quand des metadata existent.
    if metadata.get("author") or metadata.get("last_modified_by") or metadata.get("created"):
        return "Microsoft Word"
    return "Unknown"


def _has_track_changes(spans: list[dict]) -> bool:
    return any(s["is_insertion"] or s["is_deletion"] for s in spans)


def _has_comments(doc: Document) -> bool:
    """Vérifier si le doc a une part 'comments'."""
    try:
        for rel in doc.part.rels.values():
            if "comments" in (rel.reltype or ""):
                return True
    except Exception:
        pass
    return False


def _has_footnotes(doc: Document) -> bool:
    try:
        for rel in doc.part.rels.values():
            if "footnotes" in (rel.reltype or ""):
                return True
    except Exception:
        pass
    return False


def _build_doc_summary(
    docx_path: Path,
    pdf_hash: str,
    metadata: dict,
    paragraphs: list[dict],
    spans: list[dict],
    images: list[dict],
    tables: list[dict],
    sections: list[dict],
    doc: Document,
) -> dict:
    """JSON synthèse au niveau document."""
    style_counts: dict[str, int] = {}
    for p in paragraphs:
        s = p["style_name"] or "(no style)"
        style_counts[s] = style_counts.get(s, 0) + 1

    n_headings = sum(1 for p in paragraphs if p["heading_level"] is not None)
    n_list_items = sum(1 for p in paragraphs if p["is_list_item"])
    total_chars = sum(p["char_count"] for p in paragraphs)
    has_track = _has_track_changes(spans)
    has_comments = _has_comments(doc)
    has_footnotes = _has_footnotes(doc)
    has_hyperlinks = any(s["hyperlink"] for s in spans)

    return {
        "doc_hash":              pdf_hash,
        "file_size_bytes":       docx_path.stat().st_size,
        "n_paragraphs":          len(paragraphs),
        "n_spans":               len(spans),
        "n_images":              len(images),
        "n_tables":              len(tables),
        "n_sections":            len(sections),
        "n_headings":            n_headings,
        "n_list_items":          n_list_items,
        "total_char_count":      total_chars,
        # Source / metadata
        "source_tool":           _detect_source_tool(metadata),
        "metadata":              metadata,
        # Signaux structurels
        "has_toc":               n_headings > 0,
        "has_track_changes":     has_track,
        "has_comments":          has_comments,
        "has_footnotes":         has_footnotes,
        "has_hyperlinks":        has_hyperlinks,
        # Distribution des styles (top 10)
        "style_counts":          dict(sorted(style_counts.items(), key=lambda kv: -kv[1])[:10]),
        # Recommandation pipeline aval
        "recommended_strategy":  "use_native_extraction",   # un .docx est toujours extractible nativement
        # Méta-info de l'analyse
        "analyzed_at":           datetime.now().isoformat(),
        "parser_version":        "1.0.0",
    }


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║ 5. POINT D'ENTRÉE — parse_word                                             ║
# ╚════════════════════════════════════════════════════════════════════════════╝

def parse_word(docx_path) -> dict:
    """
    Parser un .docx en 5 sorties : paragraph_df, span_df, image_df, table_df,
    doc_summary. Une seule ouverture python-docx.

    Voir module docstring pour le pattern extract → modify → rebuild.
    """
    docx_path = Path(docx_path)
    doc_hash = _sha256(docx_path)
    doc = Document(str(docx_path))

    paragraphs: list[dict] = []
    spans: list[dict]      = []
    for para_idx, para in enumerate(doc.paragraphs):
        paragraphs.append(_extract_paragraph(para, para_idx))
        for run_idx, run in enumerate(para.runs):
            spans.append(_extract_span(run, para_idx, run_idx))

    # Walk tables : doc.paragraphs n'inclut PAS les paragraphs des cells
    for table_idx, table in enumerate(doc.tables):
        for row_idx, row in enumerate(table.rows):
            for col_idx, cell in enumerate(row.cells):
                for cell_para_idx, cell_para in enumerate(cell.paragraphs):
                    for cell_run_idx, cell_run in enumerate(cell_para.runs):
                        sid = (f"w_t_{table_idx}_{row_idx}_{col_idx}"
                               f"_{cell_para_idx}_{cell_run_idx}")
                        spans.append(_extract_span(
                            cell_run,
                            para_idx=cell_para_idx,
                            run_idx=cell_run_idx,
                            span_id=sid,
                            in_table=True,
                            table_pos=(table_idx, row_idx, col_idx),
                        ))

    images = [_extract_image(s, i) for i, s in enumerate(doc.inline_shapes)]
    tables = [_extract_table(t, i) for i, t in enumerate(doc.tables)]
    sections = [_extract_section(s, i) for i, s in enumerate(doc.sections)]
    metadata = _extract_doc_metadata(doc)

    paragraph_df = pd.DataFrame(paragraphs)
    span_df      = pd.DataFrame(spans)
    image_df     = pd.DataFrame(images)
    table_df     = pd.DataFrame([{k: v for k, v in t.items() if k != "cells"} for t in tables])
    # PK uniforme
    for df in (paragraph_df, span_df, image_df, table_df):
        if not df.empty:
            df.insert(0, "doc_hash", doc_hash)

    doc_summary = _build_doc_summary(
        docx_path, doc_hash, metadata, paragraphs, spans, images, tables, sections, doc,
    )

    return {
        "paragraph_df": paragraph_df,
        "span_df":      span_df,
        "image_df":     image_df,
        "table_df":     table_df,
        "doc_summary":  doc_summary,
        # Bonus utiles : sections détaillées + cellules de tables (pas dans les df)
        "sections":     sections,
        "tables_raw":   tables,
    }


# ╔════════════════════════════════════════════════════════════════════════════╗
# ║ 6. CLI minimal                                                             ║
# ╚════════════════════════════════════════════════════════════════════════════╝

if __name__ == "__main__":
    import json
    import sys

    if len(sys.argv) != 2:
        print("Usage: python parse_word.py <docx_path>", file=sys.stderr)
        sys.exit(1)

    result = parse_word(sys.argv[1])
    s = result["doc_summary"]
    print(json.dumps(s, indent=2, ensure_ascii=False, default=str))
    print()
    print(f"paragraphs : {len(result['paragraph_df'])}")
    print(f"spans      : {len(result['span_df'])}")
    print(f"images     : {len(result['image_df'])}")
    print(f"tables     : {len(result['table_df'])}")
