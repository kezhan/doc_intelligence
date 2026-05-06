"""
TODO-011 — Native Word parsing (XML, no PDF conversion)
TODO-012 — Word + PDF consolidation (TOC from Word, annotations from PDF)
"""

from __future__ import annotations

import re
import xml.etree.ElementTree as ET
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.oxml.ns import qn


# ── TODO-011 ─────────────────────────────────────────────────────────────────

@dataclass
class WordParseResult:
    """
    Output contract for TODO-011.

    df        : standardized DataFrame (one row per paragraph/span)
    toc       : list of {level, title, paragraph_index} entries
    spans     : list of {span_id, text, style, bold, italic, color, ...}
    tables    : list of DataFrames (one per native table)
    """
    df: pd.DataFrame
    toc: list[dict[str, Any]]
    spans: list[dict[str, Any]]
    tables: list[pd.DataFrame]


def parse_word(docx_path: str | Path) -> WordParseResult:
    """
    TODO-011 — Parse a .docx natively from its XML structure.

    Input : path to a .docx file
    Output: WordParseResult with standardised DataFrame + TOC + spans + tables
    """
    doc = Document(str(docx_path))

    rows: list[dict[str, Any]] = []
    spans: list[dict[str, Any]] = []
    toc: list[dict[str, Any]] = []
    span_counter = 0

    for para_idx, para in enumerate(doc.paragraphs):
        level = _heading_level(para)
        if level is not None:
            toc.append({
                "level": level,
                "title": para.text.strip(),
                "paragraph_index": para_idx,
            })

        para_spans: list[str] = []
        for run in para.runs:
            text = run.text
            if not text:
                continue
            span_id = f"w{span_counter}"
            color = _run_color(run)
            spans.append({
                "span_id": span_id,
                "paragraph_index": para_idx,
                "text": text,
                "font": run.font.name or "",
                "size": run.font.size.pt if run.font.size else None,
                "bold": bool(run.bold),
                "italic": bool(run.italic),
                "underline": bool(run.underline),
                "color": color,
                "style": para.style.name if para.style else "",
            })
            para_spans.append(span_id)
            span_counter += 1

        rows.append({
            "paragraph_index": para_idx,
            "text": para.text,
            "style": para.style.name if para.style else "",
            "heading_level": level,
            "span_ids": ",".join(para_spans),
        })

    df = pd.DataFrame(rows)
    tables = [_docx_table_to_df(t) for t in doc.tables]

    return WordParseResult(df=df, toc=toc, spans=spans, tables=tables)


# ── helpers ───────────────────────────────────────────────────────────────────

_HEADING_RE = re.compile(r"^Heading\s+(\d+)$", re.IGNORECASE)


def _heading_level(para: Any) -> int | None:
    style_name = para.style.name if para.style else ""
    m = _HEADING_RE.match(style_name)
    if m:
        return int(m.group(1))
    if style_name.lower() == "title":
        return 0
    return None


def _run_color(run: Any) -> str | None:
    if run.font.color and run.font.color.rgb:
        return f"#{run.font.color.rgb}"
    return None


def _docx_table_to_df(table: Any) -> pd.DataFrame:
    data = [[cell.text for cell in row.cells] for row in table.rows]
    if not data:
        return pd.DataFrame()
    header = data[0]
    return pd.DataFrame(data[1:], columns=header)
